import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def markdown_to_docx(md_path, docx_path):
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_table = False
    table_data = []

    def format_runs(p, text):
        # Handle bold **text**
        parts = text.split('**')
        for i, part in enumerate(parts):
            run = p.add_run(part)
            if i % 2 == 1:
                run.bold = True

    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('|') and line.endswith('|'):
            in_table = True
            # ignore separator line
            if '---' in line:
                continue
            cells = [c.strip() for c in line.split('|')[1:-1]]
            table_data.append(cells)
            continue
        elif in_table:
            # We just finished a table
            in_table = False
            if table_data:
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Table Grid'
                for r_idx, row_data in enumerate(table_data):
                    row_cells = table.rows[r_idx].cells
                    for c_idx, cell_data in enumerate(row_data):
                        if c_idx < len(row_cells):
                            row_cells[c_idx].text = cell_data
                table_data = []

        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
        elif line.startswith('* ') or line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            format_runs(p, line[2:])
        elif re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line)
            p = doc.add_paragraph(style='List Number')
            format_runs(p, text)
        elif line == '---':
            doc.add_page_break()
        elif line.startswith('```'):
            continue # ignore code blocks wrappers
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            format_runs(p, line)

    doc.save(docx_path)

if __name__ == '__main__':
    md_path = r"C:\Users\selva\.gemini\antigravity-ide\brain\de195af2-7bb4-496d-a627-f9bf59595f98\project_report.md"
    docx_path = r"Project_Report.docx"
    markdown_to_docx(md_path, docx_path)
    print(f"Created {docx_path}")
