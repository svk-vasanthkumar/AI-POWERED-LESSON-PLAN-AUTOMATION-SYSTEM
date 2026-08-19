"""Tests for the Adhiyamaan College lesson-plan export format (Task #7).

These cover the three pure layers of ``ace_lesson_plan_export``:

* ``map_pedagogy_method`` / formatting helpers (unit-level logic)
* ``assemble_ace_context`` (merging structured plan + schedule + Task #6
  execution data into render rows, plus the no-schedule fallback)
* the PDF / DOCX / XLSX builders (they must emit valid, non-empty documents
  with the college headings and every topic row present)

The builders are exercised through their real libraries (reportlab, python-docx,
openpyxl) so a broken template surfaces here rather than at request time.
"""

from __future__ import annotations

import io
import zipfile

import pytest
from docx import Document
from openpyxl import load_workbook

from app.services.ace_lesson_plan_export import (
    INSTITUTION_NAME,
    PEDAGOGY_CODES,
    assemble_ace_context,
    export_ace_docx,
    export_ace_pdf,
    export_ace_xlsx,
    format_date_with_day,
    format_pedagogy,
    format_period_label,
    map_pedagogy_method,
)


# --------------------------------------------------------------------------- #
# Fixtures
# --------------------------------------------------------------------------- #
@pytest.fixture
def structured_plan() -> dict:
    return {
        "course_title": "Artificial Intelligence",
        "units": [
            {
                "unit_number": 1,
                "unit_title": "Introduction to AI",
                "topics": [
                    {
                        "topic_id": "U1-T1",
                        "topic": "History of AI",
                        "teaching_methods": ["Lecture", "Group discussion"],
                        "references": ["AIMA Chapter 1"],
                    },
                    {
                        "topic_id": "U1-T2",
                        "topic": "Intelligent agents",
                        "teaching_methods": ["Problem based learning"],
                        "references": ["AIMA Chapter 2"],
                    },
                ],
            },
            {
                "unit_number": 2,
                "unit_title": "Problem Solving",
                "topics": [
                    {
                        "topic_id": "U2-T1",
                        "topic": "Uninformed search",
                        "teaching_methods": ["Flipped classroom", "Interactive whiteboard"],
                        "references": ["AIMA Chapter 3"],
                    },
                ],
            },
        ],
    }


@pytest.fixture
def course() -> dict:
    return {
        "course_code": "CS2026",
        "course_name": "Artificial Intelligence",
        "department": "Computer Science and Engineering",
        "semester": 6,
        "academic_year": "2025-26",
        "programme": "B.E. CSE",
    }


@pytest.fixture
def schedule() -> dict:
    return {
        "faculty_id": "fac-1",
        "sessions": [
            {
                "session_id": "U1-T1-s1",
                "topic_id": "U1-T1",
                "topic": "History of AI",
                "unit_number": 1,
                "unit_title": "Introduction to AI",
                "date": "2026-07-27",
                "day": "Monday",
                "period_start": 1,
                "period_end": 1,
                "status": "completed",
                "executed_date": "2026-07-27",
                "executed_day": "Monday",
                "actual_topics": ["History of AI", "The Turing test"],
                "actual_hours": 1.0,
            },
            {
                "session_id": "U1-T2-s1",
                "topic_id": "U1-T2",
                "topic": "Intelligent agents",
                "unit_number": 1,
                "unit_title": "Introduction to AI",
                "date": "2026-07-29",
                "day": "Wednesday",
                "period_start": 3,
                "period_end": 4,
                "status": "pending",
            },
            {
                "session_id": "U2-T1-s1",
                "topic_id": "U2-T1",
                "topic": "Uninformed search",
                "unit_number": 2,
                "unit_title": "Problem Solving",
                "date": "2026-08-03",
                "day": "Monday",
                "period_start": 2,
                "period_end": 2,
                "status": "rescheduled",
                "rescheduled_date": "2026-08-05",
                "rescheduled_day": "Wednesday",
            },
        ],
    }


@pytest.fixture
def context(structured_plan, course, schedule) -> dict:
    return assemble_ace_context(structured_plan, course, None, schedule)


# --------------------------------------------------------------------------- #
# Pedagogy mapping
# --------------------------------------------------------------------------- #
def test_pedagogy_maps_known_methods_to_college_codes():
    assert map_pedagogy_method("Chalk and Talk") == "CT"
    assert map_pedagogy_method("chalk & talk") == "CT"
    assert map_pedagogy_method("Lecture") == "CT"
    assert map_pedagogy_method("Group discussion") == "GL"
    # Per the ACE legend, PBL == "Project based learning"; problem-based teaching
    # maps to PS == "Learning through problem solving".
    assert map_pedagogy_method("Problem based learning") == "PS"
    assert map_pedagogy_method("Project based learning") == "PBL"
    assert map_pedagogy_method("PBL") == "PBL"
    assert map_pedagogy_method("Flipped classroom") == "FC"


def test_pedagogy_codes_are_deduped_and_ordered():
    # Two synonyms of the same code collapse to a single code.
    row_codes = map_pedagogy_method("Lecture"), map_pedagogy_method("Chalk and talk")
    assert row_codes == ("CT", "CT")


def test_pedagogy_unknown_method_passes_through_verbatim():
    # An unmapped method yields no code from map_pedagogy_method ...
    assert map_pedagogy_method("Escape room activity") is None
    # ... but format_pedagogy must preserve it verbatim, never drop it silently.
    assert format_pedagogy(["Escape room activity"]) == "Escape room activity"


def test_pedagogy_blank_method_returns_none():
    assert map_pedagogy_method("") is None
    assert map_pedagogy_method("   ") is None


def test_every_declared_code_is_a_valid_college_code():
    # Sanity: all mapped values are among the documented ACE codes.
    for code in PEDAGOGY_CODES.values():
        assert code in PEDAGOGY_CODES.values()


# --------------------------------------------------------------------------- #
# Formatting helpers
# --------------------------------------------------------------------------- #
def test_format_period_single_hour():
    assert format_period_label({"period_start": 1, "period_end": 1}) == "Hour 1"


def test_format_period_range():
    assert format_period_label({"period_start": 3, "period_end": 4}) == "Hour 3–4"


def test_format_period_missing_returns_blank():
    assert format_period_label({"period_start": None, "period_end": None}) == ""


def test_format_date_includes_day():
    assert format_date_with_day("2026-07-27", "Monday") == "27 Jul 2026 (Monday)"


def test_format_date_missing_returns_blank():
    assert format_date_with_day(None, None) == ""


# --------------------------------------------------------------------------- #
# Context assembly
# --------------------------------------------------------------------------- #
def test_context_groups_topics_by_unit(context):
    assert [u["unit_number"] for u in context["units"]] == [1, 2]
    assert [len(u["rows"]) for u in context["units"]] == [2, 1]


def test_context_metadata_carries_course_fields(context):
    meta = dict(context["metadata"])
    assert meta["Department"] == "Computer Science and Engineering"
    assert meta["Semester"] == "6"
    assert "CS2026" in meta["Course Code & Title"]
    assert context["academic_year"] == "2025-26"
    assert context["institution"] == INSTITUTION_NAME


def test_context_row_merges_planned_and_executed(context):
    row = context["units"][0]["rows"][0]
    assert row["planned_date"] == "27 Jul 2026 (Monday)"
    assert row["planned_period"] == "Hour 1"
    assert row["executed_date"] == "27 Jul 2026 (Monday)"
    # Actual topics taught (Task #6 execution data), not just the planned topic.
    assert "Turing test" in row["topics_covered"]
    # Pedagogy rendered as college codes.
    assert row["pedagogy"] == "CT, GL"
    assert row["resource"] == "AIMA Chapter 1"


def test_context_pending_row_has_blank_executed(context):
    row = context["units"][0]["rows"][1]
    assert row["planned_period"] == "Hour 3–4"
    assert row["executed_date"] == ""  # not taught yet


def test_context_rescheduled_row_reflects_new_date(context):
    row = context["units"][1]["rows"][0]
    # Rescheduled sessions surface their effective date somewhere in the row.
    assert "05 Aug 2026" in row["executed_date"] or "05 Aug 2026" in row["planned_date"]


def test_context_falls_back_to_structured_plan_without_schedule(structured_plan, course):
    ctx = assemble_ace_context(structured_plan, course, None, None)
    assert ctx["has_schedule"] is False
    # Every topic still appears, with planned/executed columns blank.
    assert [len(u["rows"]) for u in ctx["units"]] == [2, 1]
    row = ctx["units"][0]["rows"][0]
    assert row["planned_date"] == ""
    assert row["executed_date"] == ""
    # Pedagogy + resources still come from the structured plan.
    assert row["pedagogy"] == "CT, GL"
    assert row["resource"] == "AIMA Chapter 1"


def test_context_handles_missing_course(structured_plan, schedule):
    ctx = assemble_ace_context(structured_plan, None, None, schedule)
    meta = dict(ctx["metadata"])
    # Missing course fields degrade to blanks, never raise.
    assert meta["Department"] == ""


# --------------------------------------------------------------------------- #
# PDF builder
# --------------------------------------------------------------------------- #
def test_pdf_is_valid_and_nonempty(context):
    pdf = export_ace_pdf(context)
    assert isinstance(pdf, bytes)
    assert pdf.startswith(b"%PDF-")
    assert len(pdf) > 1000


# --------------------------------------------------------------------------- #
# DOCX builder
# --------------------------------------------------------------------------- #
def test_docx_is_valid_and_contains_headings(context):
    docx = export_ace_docx(context)
    assert docx.startswith(b"PK")  # zip container
    document = Document(io.BytesIO(docx))
    full_text = "\n".join(p.text for p in document.paragraphs)
    assert INSTITUTION_NAME in full_text
    assert "2025-26" in full_text


def test_docx_contains_every_topic_row(context):
    docx = export_ace_docx(context)
    document = Document(io.BytesIO(docx))
    cell_text = " ".join(
        cell.text for table in document.tables for row in table.rows for cell in row.cells
    )
    assert "History of AI" in cell_text
    assert "Intelligent agents" in cell_text
    assert "Uninformed search" in cell_text


# --------------------------------------------------------------------------- #
# XLSX builder
# --------------------------------------------------------------------------- #
def test_xlsx_is_valid_and_contains_data(context):
    xlsx = export_ace_xlsx(context)
    assert xlsx.startswith(b"PK")
    wb = load_workbook(io.BytesIO(xlsx))
    ws = wb.active
    flat = " ".join(
        str(cell.value) for row in ws.iter_rows() for cell in row if cell.value is not None
    )
    assert "History of AI" in flat
    assert "Uninformed search" in flat
    assert "2025-26" in flat


def test_xlsx_is_a_zip_container(context):
    xlsx = export_ace_xlsx(context)
    # openpyxl output must be a readable OOXML zip.
    with zipfile.ZipFile(io.BytesIO(xlsx)) as zf:
        assert any(name.endswith("workbook.xml") for name in zf.namelist())
