"""Tests for the extraction flow: native PDF text, DOCX, and the scanned-PDF
OCR fallback.

Real PDFs/DOCX are built with PyMuPDF and python-docx (already project deps).
The OCR engine is always mocked — no PaddleOCR/OpenCV is required — and page
rendering is stubbed so tests stay fast and hermetic.

Run:  python -m pytest backend/tests/test_text_extraction.py -q
"""

import glob
import os

import fitz  # PyMuPDF
import pytest
from docx import Document
from app.services import ocr_service
from app.services import text_extraction_service as tes
from app.services.text_extraction_service import (
    DocumentExtractionError,
    DocumentOCRProcessingError,
    DocumentOCRUnavailableError,
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_text_pdf(path: str, text: str) -> None:
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), text)
    doc.save(path)
    doc.close()


def _make_scanned_pdf(path: str, pages: int = 1) -> None:
    """A PDF with blank pages (no selectable text) — stands in for a scan."""
    doc = fitz.open()
    for _ in range(pages):
        doc.new_page()
    doc.save(path)
    doc.close()


def _make_docx(path: str, paragraphs: list[str]) -> None:
    doc = Document()
    for para in paragraphs:
        doc.add_paragraph(para)
    doc.save(path)


def _make_table_docx(
    path: str,
    tables: list[list[list[str]]],
    paragraphs_before: list[str] | None = None,
    paragraphs_after: list[str] | None = None,
) -> None:
    """Build a DOCX containing one or more tables (and optional paragraphs).

    ``tables`` is a list of tables; each table is a list of rows; each row is a
    list of cell strings.
    """
    doc = Document()
    for para in paragraphs_before or []:
        doc.add_paragraph(para)
    for table_rows in tables:
        rows = len(table_rows)
        cols = max(len(r) for r in table_rows)
        table = doc.add_table(rows=rows, cols=cols)
        for r, row in enumerate(table_rows):
            for c, value in enumerate(row):
                table.rows[r].cells[c].text = value
    for para in paragraphs_after or []:
        doc.add_paragraph(para)
    doc.save(path)


# ---------------------------------------------------------------------------
# 1. Normal text PDF -> extraction_method == "text"
# ---------------------------------------------------------------------------

def test_text_pdf_uses_native_extraction(tmp_path):
    pdf = str(tmp_path / "syllabus.pdf")
    _make_text_pdf(pdf, "Unit 1 Introduction to Algorithms and Data Structures")

    text, method = tes.extract_text_with_method(pdf)

    assert method == "text"
    assert "Algorithms" in text


# ---------------------------------------------------------------------------
# 2. Scanned PDF -> extraction_method == "ocr"
# ---------------------------------------------------------------------------

def test_scanned_pdf_uses_ocr(tmp_path, monkeypatch):
    pdf = str(tmp_path / "scanned.pdf")
    _make_scanned_pdf(pdf, pages=1)

    monkeypatch.setattr(tes, "_render_page_to_image", lambda page: object())
    monkeypatch.setattr(
        ocr_service, "run_ocr",
        lambda image, **kwargs: "Unit 1 Introduction recognised via OCR",
    )

    text, method = tes.extract_text_with_method(pdf)

    assert method == "ocr"
    assert "recognised via OCR" in text


# ---------------------------------------------------------------------------
# 3. Multi-page scanned PDF -> page order preserved
# ---------------------------------------------------------------------------

def test_multipage_scanned_pdf_preserves_order(tmp_path, monkeypatch):
    pdf = str(tmp_path / "scanned_multi.pdf")
    _make_scanned_pdf(pdf, pages=3)

    outputs = iter([
        "Unit 1 alpha content",
        "Unit 1 beta content",
        "Unit 2 gamma content",
    ])
    monkeypatch.setattr(tes, "_render_page_to_image", lambda page: object())
    monkeypatch.setattr(ocr_service, "run_ocr", lambda image, **kwargs: next(outputs))

    text, method = tes.extract_text_with_method(pdf)

    assert method == "ocr"
    assert text.index("alpha") < text.index("beta") < text.index("gamma")


# ---------------------------------------------------------------------------
# 4. DOCX -> python-docx extraction (never OCR)
# ---------------------------------------------------------------------------

def test_docx_uses_python_docx(tmp_path, monkeypatch):
    docx_path = str(tmp_path / "syllabus.docx")
    _make_docx(docx_path, ["Unit 1 Introduction", "Topic A", "Topic B"])

    # Guard: OCR must never be invoked for DOCX.
    def _fail(*args, **kwargs):
        raise AssertionError("OCR must not run for DOCX")

    monkeypatch.setattr(ocr_service, "run_ocr", _fail)

    text, method = tes.extract_text_with_method(docx_path)

    assert method == "text"
    assert "Unit 1 Introduction" in text
    assert "Topic B" in text


# ---------------------------------------------------------------------------
# 4b. DOCX table extraction (paragraphs + table cells, in document order)
# ---------------------------------------------------------------------------

def test_docx_paragraph_only(tmp_path):
    docx_path = str(tmp_path / "paras.docx")
    _make_docx(docx_path, ["Unit 1 Introduction to Algorithms", "Sorting and Searching"])

    text, method = tes.extract_text_with_method(docx_path)

    assert method == "text"
    assert "Introduction to Algorithms" in text
    assert "Sorting and Searching" in text


def test_docx_table_only(tmp_path):
    # A syllabus laid out entirely as a table must still produce meaningful text.
    docx_path = str(tmp_path / "table.docx")
    _make_table_docx(
        docx_path,
        tables=[[
            ["Unit", "Topic", "Hours"],
            ["1", "Introduction to Databases", "6"],
            ["2", "Normalization and Indexing", "8"],
        ]],
    )

    text, method = tes.extract_text_with_method(docx_path)

    assert method == "text"
    assert "Introduction to Databases" in text
    assert "Normalization and Indexing" in text
    assert "Hours" in text


def test_docx_paragraph_plus_table_preserves_order(tmp_path):
    docx_path = str(tmp_path / "mixed.docx")
    _make_table_docx(
        docx_path,
        tables=[[["Topic", "Weeks"], ["Graphs and Trees", "3"]]],
        paragraphs_before=["Course Overview Section"],
        paragraphs_after=["Assessment Pattern Details"],
    )

    text, method = tes.extract_text_with_method(docx_path)

    assert method == "text"
    # Paragraph content and table cell content both present, in document order.
    assert (
        text.index("Course Overview Section")
        < text.index("Graphs and Trees")
        < text.index("Assessment Pattern Details")
    )


def test_docx_multiple_tables(tmp_path):
    docx_path = str(tmp_path / "multi_table.docx")
    _make_table_docx(
        docx_path,
        tables=[
            [["Unit 1", "Foundations of Computing"]],
            [["Unit 2", "Operating Systems Concepts"]],
        ],
    )

    text, _ = tes.extract_text_with_method(docx_path)

    assert "Foundations of Computing" in text
    assert "Operating Systems Concepts" in text
    assert text.index("Foundations") < text.index("Operating Systems")


def test_docx_does_not_duplicate_table_text(tmp_path):
    docx_path = str(tmp_path / "no_dupe.docx")
    _make_table_docx(docx_path, tables=[[["Distributed Systems Architecture", "12"]]])

    text, _ = tes.extract_text_with_method(docx_path)

    # The cell text must appear exactly once, never duplicated by the walker.
    assert text.count("Distributed Systems Architecture") == 1


def test_empty_docx_returns_controlled_error(tmp_path):
    docx_path = str(tmp_path / "empty.docx")
    _make_docx(docx_path, [])  # no paragraphs, no tables

    with pytest.raises(DocumentExtractionError) as exc:
        tes.extract_text_with_method(docx_path)

    assert "readable text" in str(exc.value).lower()


def test_blank_paragraphs_docx_returns_controlled_error(tmp_path):
    docx_path = str(tmp_path / "blank.docx")
    _make_docx(docx_path, ["", "   ", "\t"])  # only whitespace content

    with pytest.raises(DocumentExtractionError) as exc:
        tes.extract_text_with_method(docx_path)



# ---------------------------------------------------------------------------
# 5. Corrupt / empty PDF -> controlled 400
# ---------------------------------------------------------------------------

def test_corrupt_pdf_returns_controlled_error(tmp_path):
    pdf = str(tmp_path / "corrupt.pdf")
    with open(pdf, "wb") as handle:
        handle.write(b"this is not a valid pdf payload")

    with pytest.raises(DocumentExtractionError) as exc:
        tes.extract_text_with_method(pdf)



def test_no_text_and_no_ocr_returns_empty_error(tmp_path, monkeypatch):
    pdf = str(tmp_path / "blank.pdf")
    _make_scanned_pdf(pdf, pages=1)

    monkeypatch.setattr(tes, "_render_page_to_image", lambda page: object())
    monkeypatch.setattr(ocr_service, "run_ocr", lambda image, **kwargs: "")

    with pytest.raises(DocumentExtractionError) as exc:
        tes.extract_text_with_method(pdf)

    assert "readable text" in str(exc.value).lower()


# ---------------------------------------------------------------------------
# 6. OCR failure -> controlled error (no stack trace leaks)
# ---------------------------------------------------------------------------

def test_ocr_processing_failure_maps_to_400(tmp_path, monkeypatch):
    pdf = str(tmp_path / "scanned_fail.pdf")
    _make_scanned_pdf(pdf, pages=1)

    def _boom(image, **kwargs):
        raise ocr_service.OCRProcessingError("internal detail must stay hidden")

    monkeypatch.setattr(tes, "_render_page_to_image", lambda page: object())
    monkeypatch.setattr(ocr_service, "run_ocr", _boom)

    with pytest.raises(DocumentOCRProcessingError) as exc:
        tes.extract_text_with_method(pdf)

    assert "internal detail" not in str(exc.value)


def test_ocr_unavailable_maps_to_503(tmp_path, monkeypatch):
    pdf = str(tmp_path / "scanned_unavail.pdf")
    _make_scanned_pdf(pdf, pages=1)

    def _unavailable(image, **kwargs):
        raise ocr_service.OCRUnavailableError("engine missing")

    monkeypatch.setattr(tes, "_render_page_to_image", lambda page: object())
    monkeypatch.setattr(ocr_service, "run_ocr", _unavailable)

    with pytest.raises(DocumentExtractionError) as exc:
        tes.extract_text_with_method(pdf)



# ---------------------------------------------------------------------------
# 7. Temporary OCR images are cleaned up (nothing written to disk)
# ---------------------------------------------------------------------------

def test_ocr_leaves_no_image_artifacts(tmp_path, monkeypatch):
    pdf = str(tmp_path / "scanned_clean.pdf")
    _make_scanned_pdf(pdf, pages=2)

    monkeypatch.setattr(tes, "_render_page_to_image", lambda page: object())
    monkeypatch.setattr(ocr_service, "run_ocr", lambda image, **kwargs: "Unit content here")

    tes.extract_text_with_method(pdf)

    # In-memory rendering must never drop intermediate images to disk.
    images = glob.glob(os.path.join(str(tmp_path), "**", "*.png"), recursive=True)
    images += glob.glob(os.path.join(str(tmp_path), "**", "*.jpg"), recursive=True)
    assert images == []

    uploads = os.path.abspath(
        os.path.join(os.path.dirname(__file__), "..", "app", "uploads")
    )
    if os.path.isdir(uploads):
        leftover = glob.glob(os.path.join(uploads, "*.png"))
        leftover += glob.glob(os.path.join(uploads, "*.jpg"))
        assert leftover == []


# ---------------------------------------------------------------------------
# 8. Backward-compatible extract_text() wrapper still returns a plain string
# ---------------------------------------------------------------------------

def test_extract_text_wrapper_returns_string(tmp_path):
    pdf = str(tmp_path / "wrap.pdf")
    _make_text_pdf(pdf, "Unit 1 Introduction to Software Engineering Principles")

    result = tes.extract_text(pdf)
    assert isinstance(result, str)
    assert "Software Engineering" in result


if __name__ == "__main__":
    raise SystemExit(pytest.main([__file__, "-q"]))
