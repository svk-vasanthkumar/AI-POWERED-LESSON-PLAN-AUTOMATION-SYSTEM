"""Tests for the deterministic export system (Phase 12).

Covers the pure document builders, the async fetch/dispatch layer (with an
in-memory Mongo via ``mongomock_motor``), and the HTTP endpoints (JWT required,
RBAC intact, correct content types, controlled error codes).

Run: python -m pytest backend/tests/test_export_service.py -q
"""

import asyncio
import io
import os
import sys
from datetime import datetime, UTC

sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))

import pytest
from bson import ObjectId
from fastapi import FastAPI, HTTPException
from fastapi.testclient import TestClient
from mongomock_motor import AsyncMongoMockClient

import app.database.mongodb as mongodb
from app.api.v1.lesson_plan import router as lesson_router
from app.api.v1.scheduler import router as scheduler_router
from app.auth.jwt import create_access_token
from app.services import export_service as ex


# ---------------------------------------------------------------------------
# Fixtures & sample data
# ---------------------------------------------------------------------------


@pytest.fixture()
def db(monkeypatch):
    client = AsyncMongoMockClient()
    database = client["test_db"]
    monkeypatch.setattr(mongodb, "database", database)
    return database


def _structured_plan():
    return {
        "course_title": "Artificial Intelligence",
        "course_objectives": ["Understand search", "Apply reasoning"],
        "learning_outcomes": [
            {"outcome_id": "CO1", "description": "Explain AI concepts", "bloom_level": "Understand"},
            {"outcome_id": "CO2", "description": "Apply search algorithms", "bloom_level": "Apply"},
        ],
        "units": [
            {
                "unit_number": 1,
                "unit_title": "Introduction to AI",
                "topics": [
                    {
                        "topic_id": "U1-T1",
                        "topic": "History and foundations of Artificial Intelligence and intelligent agents",
                        "subtopics": ["Turing test", "Rational agents"],
                        "estimated_hours": 2,
                        "bloom_level": "Understand",
                        "learning_outcomes": ["CO1"],
                        "teaching_methods": ["Lecture", "Discussion"],
                        "assessment_methods": ["Quiz"],
                        "references": ["AIMA Ch.1"],
                    },
                    {
                        "topic_id": "U1-T2",
                        "topic": "Problem solving by search",
                        "subtopics": ["BFS", "DFS", "A*"],
                        "estimated_hours": 3,
                        "bloom_level": "Apply",
                        "learning_outcomes": ["CO2"],
                        "teaching_methods": ["Lecture"],
                        "assessment_methods": ["Assignment"],
                        "references": ["AIMA Ch.3"],
                    },
                ],
            }
        ],
        "overall_teaching_methods": ["Lecture", "Lab"],
        "overall_assessment_methods": ["Internal exams", "End semester"],
        "references": ["Russell & Norvig, AIMA", "Poole & Mackworth"],
    }


async def _seed_lesson_plan(db, *, structured=True):
    now = datetime.now(UTC)
    course_id = ObjectId()
    await db.courses.insert_one({
        "_id": course_id,
        "course_code": "CS2026",
        "course_name": "Artificial Intelligence",
        "department": "CSE",
        "semester": 6,
        "credits": 4,
        "faculty_id": "FAC001",
        "created_at": now,
    })
    lesson_id = ObjectId()
    await db.lesson_plans.insert_one({
        "_id": lesson_id,
        "course_id": course_id,
        "syllabus_id": ObjectId(),
        "lesson_plan": "Topic A\nTopic B",
        "structured_plan": _structured_plan() if structured else None,
        "created_at": now,
    })
    return str(lesson_id), str(course_id)


async def _seed_schedule(db, *, with_sessions=True):
    now = datetime.now(UTC)
    course_id = ObjectId()
    faculty_id = ObjectId()
    await db.courses.insert_one({
        "_id": course_id, "course_code": "CS2026",
        "course_name": "Artificial Intelligence", "semester": 6,
        "faculty_id": faculty_id, "created_at": now,
    })
    await db.faculty.insert_one({
        "_id": faculty_id, "faculty_id": "FAC001", "name": "Dr. Ada Lovelace",
        "email": "ada@example.com", "department": "CSE", "designation": "Professor",
        "created_at": now,
    })
    sessions = []
    if with_sessions:
        sessions = [
            {
                "topic_id": "U1-T1", "topic": "History and foundations of AI",
                "unit_number": 1, "unit_title": "Introduction to AI",
                "date": "2026-07-27", "day": "Monday",
                "start_time": "09:00", "end_time": "10:00",
                "duration_hours": 1.0, "status": "pending",
            },
            {
                "topic_id": "U1-T2", "topic": "Problem solving by search",
                "unit_number": 1, "unit_title": "Introduction to AI",
                "date": "2026-07-29", "day": "Wednesday",
                "start_time": "11:00", "end_time": "12:00",
                "duration_hours": 1.0, "status": "pending",
            },
        ]
    await db.generated_schedules.insert_one({
        "_id": ObjectId(),
        "course_id": course_id,
        "faculty_id": faculty_id,
        "lesson_plan_id": ObjectId(),
        "calendar_id": ObjectId(),
        "timetable_id": ObjectId(),
        "sessions": sessions,
        "total_hours": 2.0,
        "version": 1,
        "active": True,
        "status": "generated",
        "created_at": now,
        "updated_at": now,
    })
    return str(course_id)


# ---------------------------------------------------------------------------
# File-signature helpers
# ---------------------------------------------------------------------------


def _is_pdf(content: bytes) -> bool:
    return content[:5] == b"%PDF-"


def _is_zip(content: bytes) -> bool:
    # DOCX and XLSX are both ZIP (OOXML) containers.
    return content[:2] == b"PK"


# ---------------------------------------------------------------------------
# 1-6, 11, 14-15: pure builders produce valid, non-empty, openable documents
# ---------------------------------------------------------------------------


def test_lesson_plan_pdf_builder():
    content = ex.export_lesson_plan_pdf(_structured_plan(), {"course_code": "CS2026"})
    assert content and _is_pdf(content)
    assert len(content) > 1000  # non-trivial document


def test_lesson_plan_docx_builder_opens():
    from docx import Document

    content = ex.export_lesson_plan_docx(_structured_plan(), {"course_code": "CS2026"})
    assert content and _is_zip(content)
    doc = Document(io.BytesIO(content))  # opens successfully
    text = "\n".join(p.text for p in doc.paragraphs)
    assert ex.INSTITUTION in text


def test_lesson_plan_xlsx_builder_sheets_and_rows():
    from openpyxl import load_workbook

    content = ex.export_lesson_plan_xlsx(_structured_plan(), {"course_code": "CS2026"})
    assert content and _is_zip(content)
    wb = load_workbook(io.BytesIO(content))
    assert wb.sheetnames == ["Lesson Plan", "Learning Outcomes", "References"]
    ws = wb["Lesson Plan"]
    # header + 2 topic rows
    assert ws.max_row == 3
    assert ws.cell(row=1, column=1).value == "Unit"
    assert ws["D2"].value == _structured_plan()["units"][0]["topics"][0]["topic"]
    assert wb["Learning Outcomes"].max_row == 3  # header + 2 outcomes


def test_schedule_pdf_builder():
    plan = None
    content = ex.export_schedule_pdf(
        {"sessions": _sched_sessions(), "total_hours": 2.0, "version": 1, "faculty_id": "x"},
        {"course_name": "AI"},
        {"name": "Dr. Ada Lovelace"},
    )
    assert content and _is_pdf(content)


def test_schedule_docx_builder_opens():
    from docx import Document

    content = ex.export_schedule_docx(
        {"sessions": _sched_sessions(), "total_hours": 2.0, "version": 1},
        {"course_name": "AI"},
        {"name": "Dr. Ada Lovelace"},
    )
    assert content and _is_zip(content)
    Document(io.BytesIO(content))


def test_schedule_xlsx_builder_sheets_and_rows():
    from openpyxl import load_workbook

    content = ex.export_schedule_xlsx(
        {"sessions": _sched_sessions(), "total_hours": 2.0, "version": 1},
        {"course_name": "AI"},
        {"name": "Dr. Ada Lovelace"},
    )
    wb = load_workbook(io.BytesIO(content))
    assert wb.sheetnames == ["Schedule", "Workload Summary"]
    ws = wb["Schedule"]
    assert ws.max_row == 3  # header + 2 sessions
    assert ws.cell(row=1, column=1).value == "Date"
    assert ws.cell(row=2, column=10).value == "Dr. Ada Lovelace"  # faculty column


def _sched_sessions():
    return [
        {"topic_id": "U1-T1", "topic": "Intro", "unit_number": 1, "unit_title": "U1",
         "date": "2026-07-27", "day": "Monday", "start_time": "09:00",
         "end_time": "10:00", "duration_hours": 1.0, "status": "pending"},
        {"topic_id": "U1-T2", "topic": "Search", "unit_number": 1, "unit_title": "U1",
         "date": "2026-07-29", "day": "Wednesday", "start_time": "11:00",
         "end_time": "12:00", "duration_hours": 1.0, "status": "pending"},
    ]


# ---------------------------------------------------------------------------
# Safe filenames (Phase 8)
# ---------------------------------------------------------------------------


def test_safe_filename_sanitizes_input():
    name = ex.safe_filename("lesson-plan", "../../etc/passwd", extension="pdf")
    assert "/" not in name and ".." not in name
    assert name.endswith(".pdf")


# ---------------------------------------------------------------------------
# Async fetch/dispatch layer (Phases 9-11)
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_build_lesson_plan_export_all_formats(db):
    lesson_id, _ = await _seed_lesson_plan(db)
    for fmt, checker in (("pdf", _is_pdf), ("docx", _is_zip), ("xlsx", _is_zip)):
        content, filename, media_type = await ex.build_lesson_plan_export(lesson_id, fmt)
        assert checker(content)
        assert filename.endswith(f".{fmt}")
        assert media_type == ex.MEDIA_TYPES[fmt]


@pytest.mark.asyncio
async def test_lesson_plan_missing_raises_404_error(db):
    with pytest.raises(ex.LessonPlanNotFoundError):
        await ex.get_lesson_plan_for_export(str(ObjectId()))


@pytest.mark.asyncio
async def test_lesson_plan_malformed_id_raises_400(db):
    with pytest.raises(HTTPException) as exc:
        await ex.get_lesson_plan_for_export("not-an-objectid")
    assert exc.value.status_code == 400


@pytest.mark.asyncio
async def test_old_lesson_plan_without_structured_raises_422(db):
    lesson_id, _ = await _seed_lesson_plan(db, structured=False)
    with pytest.raises(ex.StructuredPlanRequiredError):
        await ex.get_lesson_plan_for_export(lesson_id)


@pytest.mark.asyncio
async def test_build_schedule_export_all_formats(db):
    course_id = await _seed_schedule(db)
    for fmt, checker in (("pdf", _is_pdf), ("docx", _is_zip), ("xlsx", _is_zip)):
        content, filename, media_type = await ex.build_schedule_export(course_id, fmt)
        assert checker(content)
        assert filename.endswith(f".{fmt}")


@pytest.mark.asyncio
async def test_schedule_missing_raises(db):
    from app.services.scheduler_service import ScheduleNotFoundError

    with pytest.raises(ScheduleNotFoundError):
        await ex.get_schedule_for_export(str(ObjectId()))


@pytest.mark.asyncio
async def test_schedule_without_sessions_raises_422(db):
    course_id = await _seed_schedule(db, with_sessions=False)
    with pytest.raises(ex.EmptyScheduleError):
        await ex.get_schedule_for_export(course_id)


# ---------------------------------------------------------------------------
# HTTP endpoints: JWT required, RBAC intact, content types, error codes
# ---------------------------------------------------------------------------


def _run(coro):
    """Run an async seeding coroutine from a synchronous test."""
    return asyncio.run(coro)


@pytest.fixture()
def client(db):
    app = FastAPI()
    app.include_router(lesson_router)
    app.include_router(scheduler_router)
    return TestClient(app)


def _auth(db, role="faculty"):
    user_id = ObjectId()
    email = f"{user_id}@example.com"
    _run(
        db.users.insert_one(
            {"_id": user_id, "name": "Test User", "email": email,
             "role": role, "department": "CSE"}
        )
    )
    token = create_access_token({"sub": str(user_id), "role": role, "email": email})
    return {"Authorization": f"Bearer {token}"}


# NOTE: these are synchronous tests. Starlette's TestClient drives the ASGI app
# on its own event loop, so calling it from inside a pytest-asyncio loop would
# clash. Seeding is done via ``_run`` on a short-lived loop; mongomock is
# in-memory and loop-agnostic, so the data persists for the request.


def test_endpoint_requires_jwt(client, db):
    lesson_id, _ = _run(_seed_lesson_plan(db))
    # No Authorization header -> 401 (proves exports are not public).
    r = client.get(f"/lesson-plan/{lesson_id}/export/pdf")
    assert r.status_code == 401


def test_lesson_plan_export_endpoints_content_types(client, db):
    lesson_id, _ = _run(_seed_lesson_plan(db))
    # Task #8 scopes exports to lesson plans the caller may access; an admin is
    # an authorized manager, so this still verifies content types end-to-end.
    for fmt in ("pdf", "docx", "xlsx"):
        r = client.get(f"/lesson-plan/{lesson_id}/export/{fmt}", headers=_auth(db, "admin"))
        assert r.status_code == 200
        assert r.headers["content-type"].split(";")[0] == ex.MEDIA_TYPES[fmt]
        assert "attachment; filename=" in r.headers["content-disposition"]
        assert len(r.content) > 500


def test_lesson_plan_export_missing_404(client, db):
    r = client.get(f"/lesson-plan/{ObjectId()}/export/pdf", headers=_auth(db))
    assert r.status_code == 404


def test_lesson_plan_export_malformed_400(client, db):
    r = client.get("/lesson-plan/not-valid/export/pdf", headers=_auth(db))
    assert r.status_code == 400


def test_lesson_plan_export_old_plan_422(client, db):
    lesson_id, _ = _run(_seed_lesson_plan(db, structured=False))
    # Authorized manager reaches the export layer, which rejects a legacy plan
    # without a structured_plan (422) rather than failing on authorization.
    r = client.get(f"/lesson-plan/{lesson_id}/export/pdf", headers=_auth(db, "admin"))
    assert r.status_code == 422
    assert "Regenerate" in r.json()["detail"]


def test_schedule_export_endpoints(client, db):
    course_id = _run(_seed_schedule(db))
    for fmt in ("pdf", "docx", "xlsx"):
        r = client.get(f"/scheduler/{course_id}/export/{fmt}", headers=_auth(db))
        assert r.status_code == 200
        assert r.headers["content-type"].split(";")[0] == ex.MEDIA_TYPES[fmt]


def test_schedule_export_missing_404(client, db):
    r = client.get(f"/scheduler/{ObjectId()}/export/pdf", headers=_auth(db))
    assert r.status_code == 404


def test_rbac_still_active_delete_forbidden_for_faculty(client, db):
    # Export routes allow any authenticated role, but existing RBAC on the
    # delete endpoint must remain intact: faculty cannot delete.
    lesson_id, _ = _run(_seed_lesson_plan(db))
    r = client.request("DELETE", f"/lesson-plan/{lesson_id}", headers=_auth(db, "faculty"))
    assert r.status_code == 403
