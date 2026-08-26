"""Deterministic export of lesson plans and schedules to PDF / DOCX / XLSX.

This module is intentionally decoupled from the API layer and from the LLM:

  * It reads ONLY existing MongoDB data (structured lesson plans generated
    earlier, and previously generated schedules). It never calls the AI model
    and never generates a schedule on demand.
  * Document rendering is fully deterministic — the same stored data always
    produces the same document bytes (aside from the "generated on" timestamp).
  * All document bytes are produced in-memory (``io.BytesIO``); nothing is
    written to disk and no user input ever touches a filesystem path.

Architecture
------------
Three layers, kept separate:

  1. Async *fetch* helpers (``get_lesson_plan_for_export`` /
     ``get_schedule_for_export``) — the only part that touches Mongo. They
     validate ids, resolve relationships, and raise the controlled export
     exceptions below.
  2. Pure *builder* functions (``export_lesson_plan_pdf`` etc.) — take plain
     dicts and return ``bytes``. No DB, no network, fully unit-testable.
  3. Thin *dispatch* helpers (``build_lesson_plan_export`` /
     ``build_schedule_export``) — map a format string to the right builder and
     return ``(bytes, filename, media_type)`` for the API layer.

Exceptions (mapped to HTTP status codes by the router):
    LessonPlanNotFoundError    -> 404
    StructuredPlanRequiredError-> 422
    ScheduleNotFoundError      -> 404   (reused from scheduler_service)
    EmptyScheduleError         -> 422
    DocumentGenerationError    -> 500   (wraps any reportlab/docx/openpyxl error)
    (malformed ObjectId is raised as HTTP 400 by ``to_object_id`` upstream)
"""

from __future__ import annotations

import asyncio
import io
import re
from datetime import datetime, UTC

from bson import ObjectId

from app.config.logger import logger
from app.database.mongodb import get_database
from app.services.ace_lesson_plan_export import ACE_BUILDERS, assemble_ace_context
from app.services.scheduler_service import ScheduleNotFoundError, get_latest_schedule
from app.utils.object_id import to_object_id

INSTITUTION = "Adhiyamaan College of Engineering"

# Media types required by the spec (Phase 8).
MEDIA_TYPES = {
    "pdf": "application/pdf",
    "docx": (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ),
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
}


# ---------------------------------------------------------------------------
# Controlled exceptions
# ---------------------------------------------------------------------------


class LessonPlanNotFoundError(Exception):
    """Requested lesson plan does not exist (-> 404)."""


class StructuredPlanRequiredError(Exception):
    """Lesson plan exists but has no structured_plan (old plan) (-> 422)."""


class EmptyScheduleError(Exception):
    """Schedule exists but contains no sessions (-> 422)."""


class DocumentGenerationError(Exception):
    """Any failure inside reportlab / python-docx / openpyxl (-> 500).

    The original library exception is logged server-side; the message here is
    safe to surface to the client.
    """


# ---------------------------------------------------------------------------
# Small shared helpers
# ---------------------------------------------------------------------------


def safe_filename(*parts: str, extension: str) -> str:
    """Build a filesystem-safe download filename from arbitrary values.

    Never uses raw user input as-is: every part is slugified to
    ``[A-Za-z0-9._-]`` so nothing can traverse or escape a path. Falls back to a
    generic stem when no usable characters remain.
    """
    cleaned: list[str] = []
    for part in parts:
        if part is None:
            continue
        slug = re.sub(r"[^A-Za-z0-9]+", "-", str(part)).strip("-")
        if slug:
            cleaned.append(slug)
    stem = "-".join(cleaned) if cleaned else "export"
    stem = stem[:120]  # keep filenames sane
    return f"{stem}.{extension}"


def _as_list(value) -> list:
    if value is None:
        return []
    if isinstance(value, (list, tuple)):
        return list(value)
    return [value]


def _join(value, sep: str = ", ") -> str:
    return sep.join(str(v) for v in _as_list(value) if str(v).strip())


def _generated_on() -> str:
    return datetime.now(UTC).strftime("%d %b %Y, %H:%M UTC")


# ---------------------------------------------------------------------------
# Fetch layer (async, DB access only)
# ---------------------------------------------------------------------------


def _stringify_ids(doc: dict) -> dict:
    result = dict(doc)
    for key, value in list(result.items()):
        if isinstance(value, ObjectId):
            result[key] = str(value)
    return result


async def get_lesson_plan_for_export(lesson_plan_id: str) -> tuple[dict, dict | None]:
    """Load a lesson plan (and its course) for export.

    Raises:
        HTTPException(400): malformed ``lesson_plan_id`` (via ``to_object_id``).
        LessonPlanNotFoundError: no such lesson plan.
        StructuredPlanRequiredError: plan exists but has no ``structured_plan``.
    """
    db = get_database()
    oid = to_object_id(lesson_plan_id, field="lesson_plan_id")

    lesson = await db.lesson_plans.find_one({"_id": oid})
    if lesson is None:
        raise LessonPlanNotFoundError("Lesson plan not found")

    structured = lesson.get("structured_plan")
    if not isinstance(structured, dict) or not structured:
        # Old lesson plans may only have the flat ``lesson_plan`` markdown text.
        raise StructuredPlanRequiredError(
            "Structured lesson plan required. Regenerate the lesson plan before exporting."
        )

    course = None
    course_id = lesson.get("course_id")
    if course_id is not None:
        try:
            course = await db.courses.find_one({"_id": ObjectId(str(course_id))})
        except Exception:
            course = None
    if course:
        course = _stringify_ids(course)

    return _stringify_ids(lesson), course


async def _resolve_faculty(db, faculty_id) -> dict | None:
    """Best-effort faculty lookup by _id or by the custom ``faculty_id`` field."""
    if faculty_id is None:
        return None
    # Try native _id match.
    try:
        found = await db.faculty.find_one({"_id": ObjectId(str(faculty_id))})
        if found:
            return _stringify_ids(found)
    except Exception:
        pass
    # Fall back to the human faculty_id string field.
    found = await db.faculty.find_one({"faculty_id": str(faculty_id)})
    return _stringify_ids(found) if found else None


async def get_schedule_for_export(course_id: str) -> tuple[dict, dict | None, dict | None]:
    """Load the latest active schedule (+ course + faculty) for export.

    Reads existing data only; never generates a schedule.

    Raises:
        HTTPException(400): malformed ``course_id`` (via ``to_object_id``).
        ScheduleNotFoundError: no schedule exists for the course.
        EmptyScheduleError: schedule exists but has no sessions.
    """
    db = get_database()
    course_oid = to_object_id(course_id, field="course_id")

    # Reuse the scheduler service's canonical "latest active" lookup. It raises
    # ScheduleNotFoundError (-> 404) when nothing exists.
    schedule = await get_latest_schedule(course_id)

    sessions = schedule.get("sessions") or []
    if not sessions:
        raise EmptyScheduleError(
            "The generated schedule contains no sessions to export."
        )

    course = await db.courses.find_one({"_id": course_oid})
    course = _stringify_ids(course) if course else None

    faculty = await _resolve_faculty(db, schedule.get("faculty_id"))

    return schedule, course, faculty


# ---------------------------------------------------------------------------
# Lesson-plan builders
# ---------------------------------------------------------------------------


def _course_info_rows(structured: dict, course: dict | None) -> list[tuple[str, str]]:
    rows: list[tuple[str, str]] = []
    title = structured.get("course_title") or (course or {}).get("course_name") or "-"
    rows.append(("Course Title", str(title)))
    if course:
        if course.get("course_code"):
            rows.append(("Course Code", str(course["course_code"])))
        if course.get("department"):
            rows.append(("Department", str(course["department"])))
        if course.get("semester") is not None:
            rows.append(("Semester", str(course["semester"])))
        if course.get("credits") is not None:
            rows.append(("Credits", str(course["credits"])))
    return rows


def export_lesson_plan_pdf(structured: dict, course: dict | None) -> bytes:
    """Render a structured lesson plan to a professional, multi-page PDF."""
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import cm
        from reportlab.platypus import (
            KeepTogether,
            Paragraph,
            SimpleDocTemplate,
            Spacer,
            Table,
            TableStyle,
        )

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            topMargin=1.6 * cm,
            bottomMargin=1.6 * cm,
            leftMargin=1.6 * cm,
            rightMargin=1.6 * cm,
            title="Lesson Plan",
        )

        styles = getSampleStyleSheet()
        h_inst = ParagraphStyle(
            "Institution", parent=styles["Title"], fontSize=16, spaceAfter=2
        )
        h_sub = ParagraphStyle(
            "SubTitle",
            parent=styles["Normal"],
            fontSize=10,
            textColor=colors.HexColor("#555555"),
            spaceAfter=10,
        )
        h2 = ParagraphStyle(
            "H2", parent=styles["Heading2"], fontSize=13, spaceBefore=12, spaceAfter=6
        )
        h3 = ParagraphStyle(
            "H3", parent=styles["Heading3"], fontSize=11, spaceBefore=8, spaceAfter=4
        )
        body = ParagraphStyle(
            "Body", parent=styles["Normal"], fontSize=9.5, leading=13
        )
        cell = ParagraphStyle("Cell", parent=body, fontSize=8.5, leading=11)
        cell_head = ParagraphStyle(
            "CellHead", parent=cell, textColor=colors.white, fontName="Helvetica-Bold"
        )

        story: list = []
        story.append(Paragraph(INSTITUTION, h_inst))
        story.append(Paragraph("Lesson Plan", h_sub))
        story.append(Paragraph(f"Generated on {_generated_on()}", h_sub))

        # Course information table.
        info_rows = _course_info_rows(structured, course)
        info_table = Table(
            [[Paragraph(f"<b>{k}</b>", cell), Paragraph(v, cell)] for k, v in info_rows],
            colWidths=[4.5 * cm, 12.5 * cm],
        )
        info_table.setStyle(
            TableStyle(
                [
                    ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#cccccc")),
                    ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#f2f2f2")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 6),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                    ("TOPPADDING", (0, 0), (-1, -1), 3),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                ]
            )
        )
        story.append(info_table)

        def bullet_section(title: str, items: list) -> None:
            items = [str(i) for i in _as_list(items) if str(i).strip()]
            if not items:
                return
            story.append(Paragraph(title, h2))
            for item in items:
                story.append(Paragraph(f"&bull;&nbsp;{item}", body))

        bullet_section("Course Objectives", structured.get("course_objectives"))

        # Learning outcomes table.
        outcomes = _as_list(structured.get("learning_outcomes"))
        if outcomes:
            story.append(Paragraph("Learning Outcomes", h2))
            data = [
                [
                    Paragraph("ID", cell_head),
                    Paragraph("Description", cell_head),
                    Paragraph("Bloom Level", cell_head),
                ]
            ]
            for o in outcomes:
                if not isinstance(o, dict):
                    continue
                data.append(
                    [
                        Paragraph(str(o.get("outcome_id", "")), cell),
                        Paragraph(str(o.get("description", "")), cell),
                        Paragraph(str(o.get("bloom_level", "")), cell),
                    ]
                )
            lo_table = Table(data, colWidths=[2.2 * cm, 11.8 * cm, 3 * cm], repeatRows=1)
            lo_table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#34495e")),
                        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#cccccc")),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f7f9fa")]),
                        ("LEFTPADDING", (0, 0), (-1, -1), 5),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                        ("TOPPADDING", (0, 0), (-1, -1), 3),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                    ]
                )
            )
            story.append(lo_table)

        # Units and topics.
        units = _as_list(structured.get("units"))
        if units:
            story.append(Paragraph("Units &amp; Topics", h2))
        for unit in units:
            if not isinstance(unit, dict):
                continue
            unit_no = unit.get("unit_number", "")
            unit_title = unit.get("unit_title", "")
            header = Paragraph(f"Unit {unit_no}: {unit_title}", h3)
            topics = _as_list(unit.get("topics"))
            data = [
                [
                    Paragraph("Topic ID", cell_head),
                    Paragraph("Topic / Subtopics", cell_head),
                    Paragraph("Hrs", cell_head),
                    Paragraph("Bloom", cell_head),
                    Paragraph("Teaching", cell_head),
                    Paragraph("Assessment", cell_head),
                ]
            ]
            for t in topics:
                if not isinstance(t, dict):
                    continue
                topic_html = f"<b>{t.get('topic', '')}</b>"
                subs = _join(t.get("subtopics"), sep="; ")
                if subs:
                    topic_html += f"<br/><font size=7 color='#666666'>{subs}</font>"
                data.append(
                    [
                        Paragraph(str(t.get("topic_id", "")), cell),
                        Paragraph(topic_html, cell),
                        Paragraph(str(t.get("estimated_hours", "")), cell),
                        Paragraph(str(t.get("bloom_level") or ""), cell),
                        Paragraph(_join(t.get("teaching_methods"), sep="; "), cell),
                        Paragraph(_join(t.get("assessment_methods"), sep="; "), cell),
                    ]
                )
            topic_table = Table(
                data,
                colWidths=[1.9 * cm, 6.1 * cm, 1.0 * cm, 2.0 * cm, 3.0 * cm, 3.0 * cm],
                repeatRows=1,
            )
            topic_table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#34495e")),
                        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#cccccc")),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f7f9fa")]),
                        ("LEFTPADDING", (0, 0), (-1, -1), 4),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                        ("TOPPADDING", (0, 0), (-1, -1), 3),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                    ]
                )
            )
            story.append(Spacer(1, 4))
            # Keep the unit header with the first rows where possible.
            story.append(KeepTogether([header]))
            story.append(topic_table)

        bullet_section(
            "Overall Teaching Methods", structured.get("overall_teaching_methods")
        )
        bullet_section(
            "Overall Assessment Methods", structured.get("overall_assessment_methods")
        )
        bullet_section("References", structured.get("references"))

        doc.build(story)
        return buffer.getvalue()
    except Exception as exc:  # noqa: BLE001 - wrap library errors safely
        logger.exception("Lesson-plan PDF generation failed")
        raise DocumentGenerationError("Failed to generate lesson plan PDF") from exc


def export_lesson_plan_docx(structured: dict, course: dict | None) -> bytes:
    """Render a structured lesson plan to a professional DOCX."""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt

        document = Document()

        heading = document.add_heading(INSTITUTION, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub = document.add_paragraph("Lesson Plan")
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        gen = document.add_paragraph(f"Generated on {_generated_on()}")
        gen.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Course information table.
        document.add_heading("Course Information", level=1)
        info_rows = _course_info_rows(structured, course)
        info_table = document.add_table(rows=0, cols=2)
        info_table.style = "Light Grid Accent 1"
        for key, value in info_rows:
            cells = info_table.add_row().cells
            cells[0].text = key
            cells[1].text = value

        def bullet_section(title: str, items: list) -> None:
            items = [str(i) for i in _as_list(items) if str(i).strip()]
            if not items:
                return
            document.add_heading(title, level=1)
            for item in items:
                document.add_paragraph(item, style="List Bullet")

        bullet_section("Course Objectives", structured.get("course_objectives"))

        outcomes = _as_list(structured.get("learning_outcomes"))
        if outcomes:
            document.add_heading("Learning Outcomes", level=1)
            table = document.add_table(rows=1, cols=3)
            table.style = "Light Grid Accent 1"
            hdr = table.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text = "ID", "Description", "Bloom Level"
            for o in outcomes:
                if not isinstance(o, dict):
                    continue
                cells = table.add_row().cells
                cells[0].text = str(o.get("outcome_id", ""))
                cells[1].text = str(o.get("description", ""))
                cells[2].text = str(o.get("bloom_level", ""))

        units = _as_list(structured.get("units"))
        if units:
            document.add_heading("Units & Topics", level=1)
        for unit in units:
            if not isinstance(unit, dict):
                continue
            document.add_heading(
                f"Unit {unit.get('unit_number', '')}: {unit.get('unit_title', '')}",
                level=2,
            )
            topics = _as_list(unit.get("topics"))
            if not topics:
                continue
            table = document.add_table(rows=1, cols=6)
            table.style = "Light Grid Accent 1"
            hdr = table.rows[0].cells
            headers = [
                "Topic ID",
                "Topic",
                "Subtopics",
                "Hrs",
                "Teaching",
                "Assessment",
            ]
            for i, name in enumerate(headers):
                hdr[i].text = name
            for t in topics:
                if not isinstance(t, dict):
                    continue
                cells = table.add_row().cells
                cells[0].text = str(t.get("topic_id", ""))
                cells[1].text = str(t.get("topic", ""))
                cells[2].text = _join(t.get("subtopics"), sep="; ")
                cells[3].text = str(t.get("estimated_hours", ""))
                cells[4].text = _join(t.get("teaching_methods"), sep="; ")
                cells[5].text = _join(t.get("assessment_methods"), sep="; ")

        bullet_section(
            "Overall Teaching Methods", structured.get("overall_teaching_methods")
        )
        bullet_section(
            "Overall Assessment Methods", structured.get("overall_assessment_methods")
        )
        bullet_section("References", structured.get("references"))

        # Base font size for readability.
        style = document.styles["Normal"]
        style.font.size = Pt(10.5)

        buffer = io.BytesIO()
        document.save(buffer)
        return buffer.getvalue()
    except Exception as exc:  # noqa: BLE001
        logger.exception("Lesson-plan DOCX generation failed")
        raise DocumentGenerationError("Failed to generate lesson plan DOCX") from exc


def export_lesson_plan_xlsx(structured: dict, course: dict | None) -> bytes:
    """Render a structured lesson plan to a multi-sheet XLSX workbook."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Alignment, Font, PatternFill
        from openpyxl.utils import get_column_letter

        wb = Workbook()

        header_fill = PatternFill("solid", fgColor="34495E")
        header_font = Font(bold=True, color="FFFFFF")
        wrap = Alignment(wrap_text=True, vertical="top")

        def style_header(ws, ncols: int) -> None:
            for col in range(1, ncols + 1):
                c = ws.cell(row=1, column=col)
                c.fill = header_fill
                c.font = header_font
                c.alignment = Alignment(vertical="center")
            ws.freeze_panes = "A2"

        def autosize(ws, widths: list[int]) -> None:
            for i, width in enumerate(widths, start=1):
                ws.column_dimensions[get_column_letter(i)].width = width

        # Sheet 1: Lesson Plan (one row per topic).
        ws1 = wb.active
        ws1.title = "Lesson Plan"
        cols1 = [
            "Unit",
            "Unit Title",
            "Topic ID",
            "Topic",
            "Subtopics",
            "Estimated Hours",
            "Bloom Level",
            "Learning Outcomes",
            "Teaching Methods",
            "Assessment Methods",
        ]
        ws1.append(cols1)
        for unit in _as_list(structured.get("units")):
            if not isinstance(unit, dict):
                continue
            unit_no = unit.get("unit_number", "")
            unit_title = unit.get("unit_title", "")
            for t in _as_list(unit.get("topics")):
                if not isinstance(t, dict):
                    continue
                ws1.append(
                    [
                        unit_no,
                        unit_title,
                        t.get("topic_id", ""),
                        t.get("topic", ""),
                        _join(t.get("subtopics"), sep="; "),
                        t.get("estimated_hours", ""),
                        t.get("bloom_level") or "",
                        _join(t.get("learning_outcomes"), sep="; "),
                        _join(t.get("teaching_methods"), sep="; "),
                        _join(t.get("assessment_methods"), sep="; "),
                    ]
                )
        style_header(ws1, len(cols1))
        autosize(ws1, [8, 24, 12, 34, 40, 14, 16, 20, 26, 26])
        for row in ws1.iter_rows(min_row=2):
            for c in row:
                c.alignment = wrap

        # Sheet 2: Learning Outcomes.
        ws2 = wb.create_sheet("Learning Outcomes")
        cols2 = ["Outcome ID", "Description", "Bloom Level"]
        ws2.append(cols2)
        for o in _as_list(structured.get("learning_outcomes")):
            if not isinstance(o, dict):
                continue
            ws2.append(
                [
                    o.get("outcome_id", ""),
                    o.get("description", ""),
                    o.get("bloom_level", ""),
                ]
            )
        style_header(ws2, len(cols2))
        autosize(ws2, [14, 70, 18])
        for row in ws2.iter_rows(min_row=2):
            for c in row:
                c.alignment = wrap

        # Sheet 3: References (+ overall methods for completeness).
        ws3 = wb.create_sheet("References")
        ws3.append(["Type", "Value"])
        for ref in _as_list(structured.get("references")):
            if str(ref).strip():
                ws3.append(["Reference", str(ref)])
        for m in _as_list(structured.get("overall_teaching_methods")):
            if str(m).strip():
                ws3.append(["Teaching Method", str(m)])
        for m in _as_list(structured.get("overall_assessment_methods")):
            if str(m).strip():
                ws3.append(["Assessment Method", str(m)])
        style_header(ws3, 2)
        autosize(ws3, [20, 80])
        for row in ws3.iter_rows(min_row=2):
            for c in row:
                c.alignment = wrap

        buffer = io.BytesIO()
        wb.save(buffer)
        return buffer.getvalue()
    except Exception as exc:  # noqa: BLE001
        logger.exception("Lesson-plan XLSX generation failed")
        raise DocumentGenerationError("Failed to generate lesson plan XLSX") from exc


# ---------------------------------------------------------------------------
# Schedule builders
# ---------------------------------------------------------------------------

_SCHEDULE_COLUMNS = [
    "Date",
    "Day",
    "Start Time",
    "End Time",
    "Unit",
    "Unit Title",
    "Topic ID",
    "Topic",
    "Duration (hrs)",
    "Faculty",
    "Status",
]


def _faculty_name(faculty: dict | None, schedule: dict) -> str:
    if faculty and faculty.get("name"):
        return str(faculty["name"])
    fid = schedule.get("faculty_id")
    return str(fid) if fid else "-"


def _schedule_row(session: dict, faculty_name: str) -> list:
    # Period-based schedules may have no clock times (period numbers are the
    # canonical representation), so fall back to the Hour label for the
    # Start/End Time columns rather than leaving them blank.
    start_display = session.get("start_time", "")
    end_display = session.get("end_time", "")
    if not start_display and session.get("period_start") is not None:
        start_display = f"Hour {session['period_start']}"
    if not end_display and session.get("period_end") is not None:
        end_display = f"Hour {session['period_end']}"
    return [
        session.get("date", ""),
        session.get("day", ""),
        start_display,
        end_display,
        session.get("unit_number", ""),
        session.get("unit_title", ""),
        session.get("topic_id", ""),
        session.get("topic", ""),
        session.get("duration_hours", ""),
        faculty_name,
        session.get("status", ""),
    ]


def export_schedule_pdf(
    schedule: dict, course: dict | None, faculty: dict | None
) -> bytes:
    """Render a generated day-wise schedule to a professional, multi-page PDF."""
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import cm
        from reportlab.platypus import (
            Paragraph,
            SimpleDocTemplate,
            Spacer,
            Table,
            TableStyle,
        )

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=landscape(A4),
            topMargin=1.4 * cm,
            bottomMargin=1.4 * cm,
            leftMargin=1.2 * cm,
            rightMargin=1.2 * cm,
            title="Schedule",
        )

        styles = getSampleStyleSheet()
        h_inst = ParagraphStyle(
            "Institution", parent=styles["Title"], fontSize=16, spaceAfter=2
        )
        h_sub = ParagraphStyle(
            "SubTitle",
            parent=styles["Normal"],
            fontSize=10,
            textColor=colors.HexColor("#555555"),
            spaceAfter=8,
        )
        cell = ParagraphStyle("Cell", parent=styles["Normal"], fontSize=8, leading=10)
        cell_head = ParagraphStyle(
            "CellHead", parent=cell, textColor=colors.white, fontName="Helvetica-Bold"
        )

        story: list = []
        story.append(Paragraph(INSTITUTION, h_inst))
        title = "Day-wise Teaching Schedule"
        if course and course.get("course_name"):
            title += f" &mdash; {course['course_name']}"
        story.append(Paragraph(title, h_sub))

        faculty_name = _faculty_name(faculty, schedule)
        meta_bits = [f"Faculty: {faculty_name}"]
        if schedule.get("version") is not None:
            meta_bits.append(f"Version: {schedule['version']}")
        if schedule.get("total_hours") is not None:
            meta_bits.append(f"Total Hours: {schedule['total_hours']}")
        meta_bits.append(f"Generated on {_generated_on()}")
        story.append(Paragraph(" &nbsp;|&nbsp; ".join(meta_bits), h_sub))

        data = [[Paragraph(c, cell_head) for c in _SCHEDULE_COLUMNS]]
        for session in schedule.get("sessions", []):
            row = _schedule_row(session, faculty_name)
            data.append([Paragraph(str(v), cell) for v in row])

        col_widths = [
            2.1 * cm,  # Date
            1.8 * cm,  # Day
            1.6 * cm,  # Start
            1.6 * cm,  # End
            1.1 * cm,  # Unit
            4.2 * cm,  # Unit title
            1.8 * cm,  # Topic ID
            5.0 * cm,  # Topic
            1.6 * cm,  # Duration
            3.0 * cm,  # Faculty
            1.8 * cm,  # Status
        ]
        table = Table(data, colWidths=col_widths, repeatRows=1)
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#34495e")),
                    ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#cccccc")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f7f9fa")]),
                    ("LEFTPADDING", (0, 0), (-1, -1), 4),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                    ("TOPPADDING", (0, 0), (-1, -1), 3),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                ]
            )
        )
        story.append(table)

        # Workload summary.
        story.append(Spacer(1, 12))
        summary = Table(
            [
                [Paragraph("<b>Workload Summary</b>", cell), Paragraph("", cell)],
                [Paragraph("Faculty", cell), Paragraph(faculty_name, cell)],
                [
                    Paragraph("Total Sessions", cell),
                    Paragraph(str(len(schedule.get("sessions", []))), cell),
                ],
                [
                    Paragraph("Total Teaching Hours", cell),
                    Paragraph(str(schedule.get("total_hours", "")), cell),
                ],
            ],
            colWidths=[5 * cm, 8 * cm],
        )
        summary.setStyle(
            TableStyle(
                [
                    ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#cccccc")),
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f2f2f2")),
                    ("SPAN", (0, 0), (1, 0)),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 5),
                    ("TOPPADDING", (0, 0), (-1, -1), 3),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                ]
            )
        )
        story.append(summary)

        doc.build(story)
        return buffer.getvalue()
    except Exception as exc:  # noqa: BLE001
        logger.exception("Schedule PDF generation failed")
        raise DocumentGenerationError("Failed to generate schedule PDF") from exc


def export_schedule_docx(
    schedule: dict, course: dict | None, faculty: dict | None
) -> bytes:
    """Render a generated day-wise schedule to a professional DOCX."""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt

        document = Document()
        heading = document.add_heading(INSTITUTION, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle = "Day-wise Teaching Schedule"
        if course and course.get("course_name"):
            subtitle += f" - {course['course_name']}"
        sub = document.add_paragraph(subtitle)
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

        faculty_name = _faculty_name(faculty, schedule)
        document.add_paragraph(
            f"Faculty: {faculty_name}    |    Version: {schedule.get('version', '')}"
            f"    |    Total Hours: {schedule.get('total_hours', '')}"
        )
        document.add_paragraph(f"Generated on {_generated_on()}")

        table = document.add_table(rows=1, cols=len(_SCHEDULE_COLUMNS))
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        for i, name in enumerate(_SCHEDULE_COLUMNS):
            hdr[i].text = name
        for session in schedule.get("sessions", []):
            row = _schedule_row(session, faculty_name)
            cells = table.add_row().cells
            for i, value in enumerate(row):
                cells[i].text = str(value)

        document.add_heading("Workload Summary", level=1)
        summary = document.add_table(rows=0, cols=2)
        summary.style = "Light Grid Accent 1"
        for key, value in [
            ("Faculty", faculty_name),
            ("Total Sessions", str(len(schedule.get("sessions", [])))),
            ("Total Teaching Hours", str(schedule.get("total_hours", ""))),
        ]:
            cells = summary.add_row().cells
            cells[0].text = key
            cells[1].text = value

        document.styles["Normal"].font.size = Pt(9.5)

        buffer = io.BytesIO()
        document.save(buffer)
        return buffer.getvalue()
    except Exception as exc:  # noqa: BLE001
        logger.exception("Schedule DOCX generation failed")
        raise DocumentGenerationError("Failed to generate schedule DOCX") from exc


def export_schedule_xlsx(
    schedule: dict, course: dict | None, faculty: dict | None
) -> bytes:
    """Render a generated day-wise schedule to an XLSX workbook."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Alignment, Font, PatternFill
        from openpyxl.utils import get_column_letter

        wb = Workbook()
        header_fill = PatternFill("solid", fgColor="34495E")
        header_font = Font(bold=True, color="FFFFFF")

        faculty_name = _faculty_name(faculty, schedule)

        ws = wb.active
        ws.title = "Schedule"
        ws.append(_SCHEDULE_COLUMNS)
        for session in schedule.get("sessions", []):
            ws.append(_schedule_row(session, faculty_name))

        for col in range(1, len(_SCHEDULE_COLUMNS) + 1):
            c = ws.cell(row=1, column=col)
            c.fill = header_fill
            c.font = header_font
        ws.freeze_panes = "A2"
        widths = [12, 11, 10, 10, 6, 22, 12, 34, 12, 22, 10]
        for i, width in enumerate(widths, start=1):
            ws.column_dimensions[get_column_letter(i)].width = width

        # Workload summary sheet.
        ws2 = wb.create_sheet("Workload Summary")
        ws2.append(["Metric", "Value"])
        for key, value in [
            ("Faculty", faculty_name),
            ("Course", (course or {}).get("course_name", "")),
            ("Version", schedule.get("version", "")),
            ("Total Sessions", len(schedule.get("sessions", []))),
            ("Total Teaching Hours", schedule.get("total_hours", "")),
        ]:
            ws2.append([key, value])
        for col in range(1, 3):
            c = ws2.cell(row=1, column=col)
            c.fill = header_fill
            c.font = header_font
        ws2.column_dimensions["A"].width = 24
        ws2.column_dimensions["B"].width = 30

        buffer = io.BytesIO()
        wb.save(buffer)
        return buffer.getvalue()
    except Exception as exc:  # noqa: BLE001
        logger.exception("Schedule XLSX generation failed")
        raise DocumentGenerationError("Failed to generate schedule XLSX") from exc


# ---------------------------------------------------------------------------
# Dispatch helpers (used by the API layer)
# ---------------------------------------------------------------------------

_LESSON_PLAN_BUILDERS = {
    "pdf": export_lesson_plan_pdf,
    "docx": export_lesson_plan_docx,
    "xlsx": export_lesson_plan_xlsx,
}

_SCHEDULE_BUILDERS = {
    "pdf": export_schedule_pdf,
    "docx": export_schedule_docx,
    "xlsx": export_schedule_xlsx,
}


async def get_ace_lesson_plan_context(lesson_plan_id: str) -> tuple[dict, dict, dict | None]:
    """Assemble the Adhiyamaan College lesson-plan render context (Task #7).

    Combines the structured plan (units/topics/pedagogy/resources), the latest
    generated schedule (planned dates/periods/order) and the Task #6 execution
    data recorded on each session, plus course/faculty metadata. Reads existing
    data only — never generates a schedule and never calls the AI model.

    The schedule is optional: a lesson plan that has not been scheduled yet (or a
    legacy plan) still exports, with the planned/executed columns left blank
    (``assemble_ace_context`` falls back to the structured plan's topics).

    Raises the same controlled errors as ``get_lesson_plan_for_export``
    (LessonPlanNotFoundError -> 404, StructuredPlanRequiredError -> 422).
    """
    lesson, course = await get_lesson_plan_for_export(lesson_plan_id)
    structured = lesson["structured_plan"]

    schedule: dict | None = None
    faculty: dict | None = None
    course_id = lesson.get("course_id")
    if course_id is not None:
        try:
            schedule = await get_latest_schedule(str(course_id))
        except ScheduleNotFoundError:
            schedule = None
        except Exception as exc:  # noqa: BLE001
            logger.exception("ACE export: schedule lookup failed")
            raise DocumentGenerationError(
                "Failed to load the generated schedule for this lesson plan."
            ) from exc
        if schedule is not None:
            try:
                faculty = await _resolve_faculty(get_database(), schedule.get("faculty_id"))
            except Exception:  # noqa: BLE001
                faculty = None

    context = assemble_ace_context(structured, course, faculty, schedule)
    return context, lesson, course


async def build_lesson_plan_export(
    lesson_plan_id: str, fmt: str
) -> tuple[bytes, str, str]:
    """Fetch + render the ACE lesson plan. Returns ``(content, filename, media_type)``.

    Task #7 replaced the previous generic lesson-plan document with the exact
    Adhiyamaan College of Engineering format. The pure per-format builders
    (``export_lesson_plan_pdf`` etc.) remain in this module for backward
    compatibility and direct reuse, but the export endpoints now render the
    college format.
    """
    builder = ACE_BUILDERS[fmt]
    context, lesson, course = await get_ace_lesson_plan_context(lesson_plan_id)

    # reportlab / python-docx / openpyxl rendering is CPU-bound and blocking.
    # Run it on a worker thread so a large export never stalls the event loop
    # for the other concurrent faculty users (Task #15).
    content = await asyncio.to_thread(builder, context)

    stem_source = (
        (course or {}).get("course_code")
        or lesson["structured_plan"].get("course_title")
        or "lesson-plan"
    )
    filename = safe_filename("lesson-plan", stem_source, extension=fmt)
    return content, filename, MEDIA_TYPES[fmt]


async def build_schedule_export(course_id: str, fmt: str) -> tuple[bytes, str, str]:
    """Fetch + render a schedule. Returns ``(content, filename, media_type)``."""
    builder = _SCHEDULE_BUILDERS[fmt]
    schedule, course, faculty = await get_schedule_for_export(course_id)

    # Offload the blocking document rendering to a worker thread (Task #15).
    content = await asyncio.to_thread(builder, schedule, course, faculty)

    stem_source = (course or {}).get("course_code") or "schedule"
    filename = safe_filename("schedule", stem_source, extension=fmt)
    return content, filename, MEDIA_TYPES[fmt]
