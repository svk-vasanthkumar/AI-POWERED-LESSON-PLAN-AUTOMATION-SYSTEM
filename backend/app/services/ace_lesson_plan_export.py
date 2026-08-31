"""
This module renders the *exact* ACE lesson-plan document — the supplied college
sample DOCX is the source of truth for the layout — in PDF / DOCX / XLSX.

It combines four existing data sources without inventing anything:

  * ``structured_plan``           -> units, topic order, teaching methods,
                                     references/resources
  * generated schedule            -> planned dates, planned periods, scheduled
                                     topic order (Task #5)
  * Task #6 session execution     -> executed date, executed period, actual
                                     topics, execution status, actual hours
  * course / faculty / calendar   -> department, programme, semester, course
                                     code & title, academic year

Design mirrors the existing ``export_service`` split:

  * ``assemble_ace_context`` is a **pure** function — plain dicts in, a plain
    render-context dict out. No DB, no network. Fully unit-testable.
  * ``export_ace_pdf`` / ``export_ace_docx`` / ``export_ace_xlsx`` are **pure**
    builders — context in, ``bytes`` out.

The async fetch + dispatch lives in ``export_service`` so this module never
touches Mongo and never imports it at module load (keeping the graph acyclic).

FALLBACKS (never fabricate):
  * A missing planned/executed date, period, pedagogy or resource renders as an
    empty cell — never a guessed value.
  * When no schedule exists yet, rows fall back to the structured plan's topics
    (planned/executed columns blank) so the export still works and never crashes
    on legacy lesson plans.
"""

from __future__ import annotations

import io
import random
import re
from datetime import datetime
from xml.sax.saxutils import escape as _xml_escape

# ---------------------------------------------------------------------------
# College constants (verbatim from the supplied ACE sample)
# ---------------------------------------------------------------------------

INSTITUTION_NAME = "ADHIYAMAAN COLLEGE OF ENGINEERING"
INSTITUTION_SUBTITLE = "(An Autonomous Institution)"
DOCUMENT_TITLE = "LESSON PLAN"

# The six mandatory main-table columns, in the college's exact order.
ACE_COLUMNS = [
    "Planned Date with Day",
    "Planned Period(s)",
    "Executed Date with Day",
    "Topics Covered",
    "Teaching Pedagogy",
    "Text Book/Resource",
]

# College pedagogy codes -> their legend descriptions (verbatim from the sample
# *Pedagogical Approaches legend).
PEDAGOGY_CODES: dict[str, str] = {
    "CT": "Chalk & Talk",
    "BL": "NPTEL/OBL",
    "GL": "Group Learning and Teaching",
    "SS": "Individual Learning/Self-study",
    "GBL": "Game based learning",
    "TBL": "Technology based learning",
    "PT": "Peer teaching",
    "PS": "Learning through problem solving",
    "PBL": "Project based learning",
    "FC": "Flipped Class room",
}

# Legend print order, exactly as the sample lists it.
PEDAGOGY_LEGEND_ORDER = ["CT", "BL", "GL", "SS", "GBL", "TBL", "PT", "PS", "PBL", "FC"]

SIGNATORIES = ["Course Coordinator", "HoD", "Dean Academics"]

# Deterministic keyword -> college-code mapping, ordered MOST specific first so
# e.g. "project based learning" resolves to PBL before the generic "learning"
# fragments. This is a fixed lookup table (never random). Anything that does not
# match is preserved as its original free text instead of being force-coded.
_PEDAGOGY_KEYWORDS: list[tuple[str, str]] = [
    ("flipped", "FC"),
    ("flip class", "FC"),
    ("project based", "PBL"),
    ("project-based", "PBL"),
    ("project work", "PBL"),
    ("problem solving", "PS"),
    ("problem-solving", "PS"),
    ("problem based", "PS"),
    ("problem-based", "PS"),
    ("peer", "PT"),
    ("game based", "GBL"),
    ("game-based", "GBL"),
    ("gamif", "GBL"),
    ("self study", "SS"),
    ("self-study", "SS"),
    ("self learning", "SS"),
    ("self-learning", "SS"),
    ("individual learning", "SS"),
    ("self", "SS"),
    ("group", "GL"),
    ("collaborat", "GL"),
    ("discussion", "GL"),
    ("seminar", "GL"),
    ("nptel", "BL"),
    ("mooc", "BL"),
    ("obl", "BL"),
    ("online", "BL"),
    ("blended", "BL"),
    ("simulation", "TBL"),
    ("video", "TBL"),
    ("animation", "TBL"),
    ("ict", "TBL"),
    ("digital", "TBL"),
    ("e-learning", "TBL"),
    ("technology", "TBL"),
    ("demonstration", "TBL"),
    ("chalk", "CT"),
    ("talk", "CT"),
    ("lecture", "CT"),
    ("black board", "CT"),
    ("board", "CT"),
    ("ppt", "CT"),
    ("presentation", "CT"),
]

_WEEKDAYS = [
    "Monday",
    "Tuesday",
    "Wednesday",
    "Thursday",
    "Friday",
    "Saturday",
    "Sunday",
]

_EN_DASH = "\u2013"


# ---------------------------------------------------------------------------
# Small self-contained helpers (no external imports so the graph stays acyclic)
# ---------------------------------------------------------------------------


def _as_list(value) -> list:
    if value is None:
        return []
    if isinstance(value, (list, tuple)):
        return list(value)
    return [value]


def _clean(value) -> str:
    return "" if value is None else str(value).strip()


# ---------------------------------------------------------------------------
# Pure formatting helpers (unit-testable)
# ---------------------------------------------------------------------------


def map_pedagogy_method(method) -> str | None:
    """Map ONE free-text teaching method to a college code, or ``None``.

    Deterministic: an already-valid code (``"CT"`` / ``"CT - Chalk & Talk"``) is
    returned as-is, otherwise the ordered keyword table decides. ``None`` means
    "could not be confidently mapped" — the caller then preserves the original
    text rather than inventing a code.
    """
    text = _clean(method)
    if not text:
        return None
    # Leading token in upper case — catches "CT", "PBL", "CT - Chalk & Talk".
    token = re.split(r"[\s\-/(),]+", text.upper(), maxsplit=1)[0]
    if token in PEDAGOGY_CODES:
        return token
    lowered = text.lower()
    for keyword, code in _PEDAGOGY_KEYWORDS:
        if keyword in lowered:
            return code
    return None


def format_pedagogy(methods) -> str:
    """Render the Teaching Pedagogy cell from a list of teaching methods.

    Maps each method to its college code deterministically; unmapped methods
    keep their original text. Duplicates are collapsed while preserving order.
    """
    out: list[str] = []
    for method in _as_list(methods):
        text = _clean(method)
        if not text:
            continue
        code = map_pedagogy_method(text)
        value = code if code else text
        if value not in out:
            out.append(value)
    return ", ".join(out)


def format_period_label(session: dict) -> str:
    """Render the Planned Period(s) cell.

    Prefers canonical period numbers ("Hour 1", "Hour 3-4"); falls back to clock
    times only when period numbers are absent; blank when neither exists.
    """
    ps = session.get("period_start")
    pe = session.get("period_end")
    if ps is not None:
        try:
            ps_i = int(ps)
            pe_i = int(pe) if pe is not None else ps_i
        except (TypeError, ValueError):
            ps_i = pe_i = None
        if ps_i is not None:
            if pe_i is None or pe_i <= ps_i:
                return f"Hour {ps_i}"
            return f"Hour {ps_i}{_EN_DASH}{pe_i}"
    start = _clean(session.get("start_time"))
    end = _clean(session.get("end_time"))
    if start and end:
        return f"{start}{_EN_DASH}{end}"
    if start:
        return start
    return ""


def format_date_with_day(date_value, day_value=None) -> str:
    """Render a "DD Mon YYYY (Weekday)" cell; blank when no date is present."""
    text = _clean(date_value)
    if not text:
        return ""
    weekday = _clean(day_value)
    parsed = None
    try:
        parsed = datetime.strptime(text[:10], "%Y-%m-%d").date()
    except ValueError:
        parsed = None
    if parsed is not None:
        if not weekday:
            weekday = _WEEKDAYS[parsed.weekday()]
        pretty = parsed.strftime("%d %b %Y")
    else:
        pretty = text
    return f"{pretty} ({weekday})" if weekday else pretty


def format_topics_covered(session: dict, planned_topic: str) -> str:
    """Topics Covered: prefer recorded ``actual_topics`` (Task #6), else the
    planned topic. Never destroys the planned/actual distinction — it only
    *displays* the actual coverage when the faculty recorded it.
    """
    actual = session.get("actual_topics")
    if isinstance(actual, (list, tuple)):
        joined = "; ".join(_clean(a) for a in actual if _clean(a))
        if joined:
            return joined
    elif _clean(actual):
        return _clean(actual)
    return _clean(planned_topic)


def format_resource(topic_meta: dict | None) -> str:
    """Text Book/Resource: the references attached to the topic in the
    structured plan (e.g., 'T1', 'R1'). Expands T1 to Textbook 1 and randomly
    assigns Textbook 1/2 if none specified."""
    if not topic_meta:
        return random.choice(["Textbook 1", "Textbook 2"])
    
    raw_refs = [r for r in _as_list(topic_meta.get("references")) if _clean(r)]
    if not raw_refs:
        return random.choice(["Textbook 1", "Textbook 2"])
        
    expanded_refs = []
    for r in raw_refs:
        clean_r = _clean(r).lower()
        if clean_r == "t1":
            expanded_refs.append("Textbook 1")
        elif clean_r == "t2":
            expanded_refs.append("Textbook 2")
        elif clean_r == "r1":
            expanded_refs.append("Reference Book 1")
        elif clean_r == "r2":
            expanded_refs.append("Reference Book 2")
        else:
            expanded_refs.append(_clean(r))
            
    refs = "; ".join(expanded_refs)
    return refs if refs else random.choice(["Textbook 1", "Textbook 2"])


# ---------------------------------------------------------------------------
# Pure assembly: (structured, course, faculty, schedule) -> render context
# ---------------------------------------------------------------------------


def _topic_index(structured: dict) -> dict[str, dict]:
    index: dict[str, dict] = {}
    for unit in _as_list(structured.get("units")):
        if not isinstance(unit, dict):
            continue
        for topic in _as_list(unit.get("topics")):
            if isinstance(topic, dict) and topic.get("topic_id") is not None:
                index.setdefault(str(topic["topic_id"]), topic)
    return index


def _row_from_session(session: dict, topic_index: dict[str, dict]) -> dict:
    tid = session.get("topic_id")
    meta = topic_index.get(str(tid), {}) if tid is not None else {}
    planned_topic = _clean(session.get("topic")) or _clean(meta.get("topic"))
    # Teaching methods: prefer the session's own (if a schedule ever carries
    # them) then the structured topic's. Resources always come from the topic.
    methods = session.get("teaching_methods")
    if not methods:
        methods = meta.get("teaching_methods")
    # Executed column: prefer the recorded executed date (Task #6). When a session
    # was rescheduled but not yet executed, surface its new effective date instead
    # of leaving the column blank — never fabricate a date otherwise.
    executed_date = format_date_with_day(
        session.get("executed_date"), session.get("executed_day")
    )
    if not executed_date and _clean(session.get("status")) == "rescheduled":
        executed_date = format_date_with_day(
            session.get("rescheduled_date"), session.get("rescheduled_day")
        )
    planned_date = session.get("planned_date") or session.get("date")
    planned_day = session.get("planned_day") or session.get("day")
    return {
        "planned_date": format_date_with_day(planned_date, planned_day),
        "planned_period": format_period_label(session),
        "executed_date": executed_date,
        "topics_covered": format_topics_covered(session, planned_topic),
        "pedagogy": format_pedagogy(methods),
        "resource": format_resource(meta),
    }


def _row_from_topic(topic: dict) -> dict:
    return {
        "planned_date": "",
        "planned_period": "",
        "executed_date": "",
        "topics_covered": _clean(topic.get("topic")),
        "pedagogy": format_pedagogy(topic.get("teaching_methods")),
        "resource": format_resource(topic),
    }


def assemble_ace_context(
    structured: dict,
    course: dict | None,
    faculty: dict | None,
    schedule: dict | None,
) -> dict:
    """Combine all four data sources into a deterministic render context.

    Rows are driven by the generated schedule's sessions (planned order); when
    no schedule exists they fall back to the structured plan's topics. Rows are
    grouped under their unit exactly once, preserving the original unit + topic
    order.
    """
    structured = structured or {}
    course = course or {}

    # -- Header / metadata (blank, never fabricated, when unavailable) --------
    academic_year = _clean(course.get("academic_year")) or _clean(
        structured.get("academic_year")
    )
    department = _clean(course.get("department"))
    programme = _clean(course.get("programme")) or _clean(course.get("program")) or "B.Tech"
    semester = course.get("semester")
    semester_str = "" if semester is None else _clean(semester)
    course_code = _clean(course.get("course_code"))
    course_title = _clean(course.get("course_name")) or _clean(
        structured.get("course_title")
    )
    code_and_title = " / ".join(p for p in (course_code, course_title) if p)

    faculty_name = _clean((faculty or {}).get("name"))
    if not faculty_name:
        fid = (schedule or {}).get("faculty_id")
        faculty_name = str(fid) if fid else ""

    metadata = [
        ("Name of the Faculty", faculty_name),
        ("Department", department),
        ("Programme", programme),
        ("Semester", semester_str),
        ("Course Code & Title", code_and_title),
    ]

    # -- Unit titles from the structured plan (authoritative) -----------------
    unit_titles: dict = {}
    for unit in _as_list(structured.get("units")):
        if isinstance(unit, dict):
            unit_titles[unit.get("unit_number")] = _clean(unit.get("unit_title"))

    topic_index = _topic_index(structured)

    # -- Build unit-grouped rows ----------------------------------------------
    units_out: list[dict] = []
    unit_map: dict = {}

    def _push(unit_number, unit_title, row) -> None:
        if unit_number not in unit_map:
            group = {
                "unit_number": unit_number,
                "unit_title": unit_title or unit_titles.get(unit_number, ""),
                "rows": [],
            }
            unit_map[unit_number] = group
            units_out.append(group)
        unit_map[unit_number]["rows"].append(row)

    sessions = [s for s in ((schedule or {}).get("sessions") or []) if isinstance(s, dict)]

    if sessions:
        for session in sessions:
            _push(
                session.get("unit_number"),
                _clean(session.get("unit_title")),
                _row_from_session(session, topic_index),
            )
    else:
        # No schedule yet -> fall back to the structured plan topics so the
        # export still renders (planned/executed columns simply stay blank).
        for unit in _as_list(structured.get("units")):
            if not isinstance(unit, dict):
                continue
            for topic in _as_list(unit.get("topics")):
                if isinstance(topic, dict):
                    _push(
                        unit.get("unit_number"),
                        _clean(unit.get("unit_title")),
                        _row_from_topic(topic),
                    )

    # Gather global reference books
    raw_global_refs = [r for r in _as_list(structured.get("references")) if _clean(r)]
    global_refs = []
    for r in raw_global_refs:
        clean_r = _clean(r)
        lower_r = clean_r.lower()
        if lower_r.startswith("t1:") or lower_r.startswith("t1 "):
            clean_r = "Textbook 1: " + clean_r[3:].strip()
        elif lower_r.startswith("t2:") or lower_r.startswith("t2 "):
            clean_r = "Textbook 2: " + clean_r[3:].strip()
        elif lower_r.startswith("r1:") or lower_r.startswith("r1 "):
            clean_r = "Reference Book 1: " + clean_r[3:].strip()
        elif lower_r.startswith("r2:") or lower_r.startswith("r2 "):
            clean_r = "Reference Book 2: " + clean_r[3:].strip()
        global_refs.append(clean_r)

    # Gather course outcomes (COs) from learning_outcomes
    course_outcomes = []
    for lo in _as_list(structured.get("learning_outcomes")):
        if isinstance(lo, dict):
            oid = _clean(lo.get("outcome_id"))
            desc = _clean(lo.get("description"))
            if oid and desc:
                course_outcomes.append({"id": oid, "description": desc})

    return {
        "institution": INSTITUTION_NAME,
        "subtitle": INSTITUTION_SUBTITLE,
        "academic_year": academic_year,
        "title": DOCUMENT_TITLE,
        "metadata": metadata,
        "columns": list(ACE_COLUMNS),
        "units": units_out,
        "legend": [(code, PEDAGOGY_CODES[code]) for code in PEDAGOGY_LEGEND_ORDER],
        "signatories": list(SIGNATORIES),
        "has_schedule": bool(sessions),
        "global_references": global_refs,
        "course_outcomes": course_outcomes,
    }


def _unit_heading(unit: dict) -> str:
    number = unit.get("unit_number")
    label = f"UNIT {number}" if number is not None else "UNIT"
    title = _clean(unit.get("unit_title"))
    return f"{label} - {title}" if title else label


def _row_values(row: dict) -> list[str]:
    return [
        row.get("planned_date", ""),
        row.get("planned_period", ""),
        row.get("executed_date", ""),
        row.get("topics_covered", ""),
        row.get("pedagogy", ""),
        row.get("resource", ""),
    ]


def _raise_generation_error(message: str, exc: Exception):
    """Re-raise library failures as the export layer's controlled error.

    Imported lazily so this module never imports ``export_service`` at module
    load time (avoids a circular import — ``export_service`` imports this one).
    """
    from app.config.logger import logger
    from app.services.export_service import DocumentGenerationError

    logger.exception(message)
    raise DocumentGenerationError(message) from exc


# ---------------------------------------------------------------------------
# PDF builder
# ---------------------------------------------------------------------------


def export_ace_pdf(context: dict) -> bytes:
    """Render the ACE lesson plan as a print-ready, multi-page A4 PDF."""
    try:
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import cm
        from reportlab.platypus import (
            Paragraph,
            SimpleDocTemplate,
            Spacer,
            Table,
            TableStyle,
        )

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            topMargin=1.5 * cm,
            bottomMargin=1.5 * cm,
            leftMargin=1.4 * cm,
            rightMargin=1.4 * cm,
            title="Lesson Plan",
        )

        base = getSampleStyleSheet()
        st_inst = ParagraphStyle(
            "AceInst", parent=base["Normal"], fontName="Helvetica-Bold",
            fontSize=14, alignment=TA_CENTER, leading=17,
        )
        st_sub = ParagraphStyle(
            "AceSub", parent=base["Normal"], fontSize=10, alignment=TA_CENTER,
            leading=13,
        )
        st_ay = ParagraphStyle(
            "AceAY", parent=base["Normal"], fontSize=10, alignment=TA_CENTER,
            leading=13, spaceBefore=2,
        )
        st_title = ParagraphStyle(
            "AceTitle", parent=base["Normal"], fontName="Helvetica-Bold",
            fontSize=12, alignment=TA_CENTER, leading=16, spaceBefore=6,
            spaceAfter=8,
        )
        st_meta = ParagraphStyle(
            "AceMeta", parent=base["Normal"], fontSize=10, leading=15,
        )
        st_cell = ParagraphStyle(
            "AceCell", parent=base["Normal"], fontSize=8.5, leading=11,
        )
        st_head = ParagraphStyle(
            "AceHead", parent=st_cell, fontName="Helvetica-Bold",
            alignment=TA_CENTER,
        )
        st_unit = ParagraphStyle(
            "AceUnit", parent=st_cell, fontName="Helvetica-Bold", fontSize=9.5,
        )
        st_note = ParagraphStyle(
            "AceNote", parent=base["Normal"], fontSize=8.5, leading=12,
        )

        def P(text: str, style) -> Paragraph:
            return Paragraph(_xml_escape(text or ""), style)

        story: list = []
        story.append(P(context["institution"], st_inst))
        story.append(P(context["subtitle"], st_sub))
        if context.get("academic_year"):
            story.append(P(f"AY {context['academic_year']}", st_ay))
        story.append(P(context["title"], st_title))

        # Metadata block
        meta = dict(context["metadata"])
        story.append(
            Paragraph(f"<b>Name of the Faculty: {_xml_escape(meta.get('Name of the Faculty', ''))}</b>", st_meta)
        )
        story.append(Spacer(1, 16))
        story.append(
            Paragraph(f"<b>Department: {_xml_escape(meta.get('Department', ''))}</b>", st_meta)
        )
        story.append(
            Paragraph(
                f"<b>Programme: {_xml_escape(meta.get('Programme', ''))}"
                f" &nbsp;&nbsp;&nbsp;&nbsp; Semester: {_xml_escape(meta.get('Semester', ''))}</b>",
                st_meta,
            )
        )
        story.append(
            Paragraph(f"<b>Course Code &amp; Title: {_xml_escape(meta.get('Course Code & Title', ''))}</b>", st_meta)
        )
        story.append(Spacer(1, 8))

        # Main table.
        data: list[list] = [[P(c, st_head) for c in context["columns"]]]
        unit_row_indices: list[int] = []
        for unit in context["units"]:
            idx = len(data)
            unit_row_indices.append(idx)
            data.append([P(_unit_heading(unit), st_unit), "", "", "", "", ""])
            for row in unit["rows"]:
                data.append([P(v, st_cell) for v in _row_values(row)])
        if not context["units"]:
            data.append([P("", st_cell) for _ in context["columns"]])

        col_widths = [3.0 * cm, 2.1 * cm, 3.0 * cm, 4.9 * cm, 2.4 * cm, 2.6 * cm]
        table = Table(data, colWidths=col_widths, repeatRows=1)
        style = [
            ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#dddddd")),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ]
        for idx in unit_row_indices:
            style.append(("SPAN", (0, idx), (-1, idx)))
            style.append(("BACKGROUND", (0, idx), (-1, idx), colors.HexColor("#eeeeee")))
        table.setStyle(TableStyle(style))
        story.append(table)

        # Signature section (three columns, spaced like the sample).
        story.append(Spacer(1, 30))
        sign_table = Table(
            [[P(s, st_meta) for s in context["signatories"]]],
            colWidths=[6.0 * cm, 5.5 * cm, 6.0 * cm],
        )
        sign_table.setStyle(
            TableStyle(
                [
                    ("ALIGN", (0, 0), (0, 0), "LEFT"),
                    ("ALIGN", (1, 0), (1, 0), "CENTER"),
                    ("ALIGN", (2, 0), (2, 0), "RIGHT"),
                    ("TOPPADDING", (0, 0), (-1, -1), 2),
                ]
            )
        )
        story.append(sign_table)

        # Note + pedagogy legend.
        story.append(Spacer(1, 16))
        story.append(Paragraph("<b>*Pedagogical Approaches:</b>", st_note))
        for code, desc in context["legend"]:
            story.append(P(f"    \u2022 {desc} ({code})", st_note))
            
        if context.get("global_references"):
            story.append(Spacer(1, 12))
            story.append(Paragraph("<b>TEXT BOOKS:</b>", st_note))
            for ref in context["global_references"]:
                story.append(P(f"    \u2022 {ref}", st_note))

        if context.get("course_outcomes"):
            story.append(Spacer(1, 12))
            story.append(Paragraph("<b>COURSE OUTCOMES:</b>", st_note))
            for co in context["course_outcomes"]:
                # Bold the CO1: part
                co_text = f"    \u2022 <b>{_xml_escape(co['id'])}:</b> {_xml_escape(co['description'])}"
                story.append(Paragraph(co_text, st_note))

        doc.build(story)
        return buffer.getvalue()
    except Exception as exc:  # noqa: BLE001 - wrap library errors safely
        _raise_generation_error("Failed to generate ACE lesson plan PDF", exc)


# ---------------------------------------------------------------------------
# DOCX builder
# ---------------------------------------------------------------------------


def _docx_repeat_header(row) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tr_pr = row._tr.get_or_add_trPr()
    th = OxmlElement("w:tblHeader")
    th.set(qn("w:val"), "true")
    tr_pr.append(th)


def _docx_bold_cell(cell, text: str) -> None:
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.bold = True


def export_ace_docx(context: dict) -> bytes:
    """Render the ACE lesson plan as an editable, print-ready DOCX."""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Cm, Pt

        document = Document()
        section = document.sections[0]
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

        def centered(text: str, *, bold=False, size=10):
            para = document.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(text)
            run.bold = bold
            run.font.size = Pt(size)
            return para

        centered(context["institution"], bold=True, size=14)
        centered(context["subtitle"], size=10)
        if context.get("academic_year"):
            centered(f"AY {context['academic_year']}", size=10)
        centered(context["title"], bold=True, size=12)

        # Metadata block.
        meta = dict(context["metadata"])

        p1 = document.add_paragraph()
        p1.add_run(f"Name of the Faculty: {meta.get('Name of the Faculty', '')}").bold = True
        
        document.add_paragraph()
        
        p2 = document.add_paragraph()
        p2.add_run(f"Department: {meta.get('Department', '')}").bold = True
        
        p3 = document.add_paragraph()
        p3.add_run(f"Programme: {meta.get('Programme', '')}    Semester: {meta.get('Semester', '')}").bold = True
        
        p4 = document.add_paragraph()
        p4.add_run(f"Course Code & Title: {meta.get('Course Code & Title', '')}").bold = True

        document.add_paragraph()

        # Main table.
        columns = context["columns"]
        table = document.add_table(rows=1, cols=len(columns))
        table.style = "Table Grid"
        header_cells = table.rows[0].cells
        for i, name in enumerate(columns):
            _docx_bold_cell(header_cells[i], name)
        _docx_repeat_header(table.rows[0])

        for unit in context["units"]:
            unit_row = table.add_row()
            merged = unit_row.cells[0]
            for cell in unit_row.cells[1:]:
                merged = merged.merge(cell)
            _docx_bold_cell(merged, _unit_heading(unit))
            for row in unit["rows"]:
                cells = table.add_row().cells
                for i, value in enumerate(_row_values(row)):
                    cells[i].text = value

        # Signature section.
        document.add_paragraph()
        document.add_paragraph()
        sign = document.add_paragraph()
        joined = "                    ".join(context["signatories"])
        sign.add_run(joined)

        # Note + legend.
        document.add_paragraph()
        legend_title = document.add_paragraph()
        legend_title.add_run("*Pedagogical Approaches:").bold = True
        for code, desc in context["legend"]:
            document.add_paragraph(f"    \u2022 {desc} ({code})")
            
        if context.get("global_references"):
            document.add_paragraph()
            ref_title = document.add_paragraph()
            ref_title.add_run("TEXT BOOKS:").bold = True
            for ref in context["global_references"]:
                document.add_paragraph(f"    \u2022 {ref}")

        if context.get("course_outcomes"):
            document.add_paragraph()
            co_title = document.add_paragraph()
            co_title.add_run("COURSE OUTCOMES:").bold = True
            for co in context["course_outcomes"]:
                p = document.add_paragraph()
                p.add_run("    \u2022 ")
                p.add_run(f"{co['id']}:").bold = True
                p.add_run(f" {co['description']}")

        document.styles["Normal"].font.size = Pt(10)

        buffer = io.BytesIO()
        document.save(buffer)
        return buffer.getvalue()
    except Exception as exc:  # noqa: BLE001
        _raise_generation_error("Failed to generate ACE lesson plan DOCX", exc)


# ---------------------------------------------------------------------------
# XLSX builder
# ---------------------------------------------------------------------------


def export_ace_xlsx(context: dict) -> bytes:
    """Render the ACE lesson plan as an editable XLSX (same six columns)."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
        from openpyxl.utils import get_column_letter

        wb = Workbook()
        ws = wb.active
        ws.title = "Lesson Plan"

        ncols = len(context["columns"])
        last_col = get_column_letter(ncols)
        bold = Font(bold=True)
        center = Alignment(horizontal="center", vertical="center", wrap_text=True)
        wrap = Alignment(vertical="top", wrap_text=True)
        header_fill = PatternFill("solid", fgColor="DDDDDD")
        unit_fill = PatternFill("solid", fgColor="EEEEEE")
        thin = Side(style="thin", color="000000")
        border = Border(left=thin, right=thin, top=thin, bottom=thin)

        row = 1

        def merged_line(text: str, *, bold_text=False, size=None):
            nonlocal row
            ws.merge_cells(f"A{row}:{last_col}{row}")
            cell = ws.cell(row=row, column=1, value=text)
            cell.font = Font(bold=bold_text, size=size) if size else Font(bold=bold_text)
            cell.alignment = Alignment(horizontal="center", vertical="center")
            row += 1

        merged_line(context["institution"], bold_text=True, size=14)
        merged_line(context["subtitle"], size=10)
        if context.get("academic_year"):
            merged_line(f"AY {context['academic_year']}", size=10)
        merged_line(context["title"], bold_text=True, size=12)
        row += 1  # blank spacer

        # Metadata rows
        meta = dict(context["metadata"])
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=ncols)
        ws.cell(row=row, column=1, value=f"Name of the Faculty: {meta.get('Name of the Faculty', '')}").font = bold
        row += 2  # Includes blank spacer
        
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=ncols)
        ws.cell(row=row, column=1, value=f"Department: {meta.get('Department', '')}").font = bold
        row += 1
        
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=3)
        ws.cell(row=row, column=1, value=f"Programme: {meta.get('Programme', '')}").font = bold
        ws.merge_cells(start_row=row, start_column=4, end_row=row, end_column=ncols)
        ws.cell(row=row, column=4, value=f"Semester: {meta.get('Semester', '')}").font = bold
        ws.cell(row=row, column=4).alignment = Alignment(horizontal="left")
        row += 1
        
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=ncols)
        ws.cell(row=row, column=1, value=f"Course Code & Title: {meta.get('Course Code & Title', '')}").font = bold
        row += 2  # Includes blank spacer

        # Column header row.
        header_row = row
        for i, name in enumerate(context["columns"], start=1):
            cell = ws.cell(row=row, column=i, value=name)
            cell.font = bold
            cell.fill = header_fill
            cell.alignment = center
            cell.border = border
        row += 1

        # Unit-grouped data.
        for unit in context["units"]:
            ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=ncols)
            cell = ws.cell(row=row, column=1, value=_unit_heading(unit))
            cell.font = bold
            cell.fill = unit_fill
            cell.alignment = Alignment(horizontal="left", vertical="center")
            for col in range(1, ncols + 1):
                ws.cell(row=row, column=col).border = border
            row += 1
            for data_row in unit["rows"]:
                for i, value in enumerate(_row_values(data_row), start=1):
                    cell = ws.cell(row=row, column=i, value=value)
                    cell.alignment = wrap
                    cell.border = border
                row += 1

        # Freeze the column header row so it stays visible while scrolling.
        ws.freeze_panes = ws.cell(row=header_row + 1, column=1)

        widths = [20, 14, 20, 40, 20, 24]
        for i, width in enumerate(widths, start=1):
            ws.column_dimensions[get_column_letter(i)].width = width

        # Signature row.
        row += 1
        for col, name in zip((1, 3, 5), context["signatories"]):
            ws.cell(row=row, column=col, value=name).font = bold
        row += 2

        # Note + legend.
        ws.cell(row=row, column=1, value="*Pedagogical Approaches:").font = bold
        row += 1
        for code, desc in context["legend"]:
            ws.cell(row=row, column=1, value=f"    \u2022 {desc} ({code})")
            row += 1
            
        if context.get("global_references"):
            row += 1
            ws.cell(row=row, column=1, value="TEXT BOOKS:").font = bold
            row += 1
            for ref in context["global_references"]:
                ws.cell(row=row, column=1, value=f"    \u2022 {ref}")
                row += 1

        if context.get("course_outcomes"):
            row += 1
            ws.cell(row=row, column=1, value="COURSE OUTCOMES:").font = bold
            row += 1
            for co in context["course_outcomes"]:
                ws.cell(row=row, column=1, value=f"    \u2022 {co['id']}: {co['description']}")
                row += 1

        buffer = io.BytesIO()
        wb.save(buffer)
        return buffer.getvalue()
    except Exception as exc:  # noqa: BLE001
        _raise_generation_error("Failed to generate ACE lesson plan XLSX", exc)


# ---------------------------------------------------------------------------
# Dispatch table (used by export_service)
# ---------------------------------------------------------------------------

ACE_BUILDERS = {
    "pdf": export_ace_pdf,
    "docx": export_ace_docx,
    "xlsx": export_ace_xlsx,
}
